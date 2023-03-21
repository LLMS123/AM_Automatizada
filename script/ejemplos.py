from docxtpl import DocxTemplate
from docxtpl import InlineImage
from docx.shared import Cm
import docx
from docx import Document

doc = Document(r'C:\Users\ANALISTAUP29\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada\Am_prueba.docx')

#doc.add_paragraph('<<Indice>>')
#doc.render({})
#doc.update_toc()

doc.add_heading("1. Intervenciones y Acciones Pedagógicas  (IAP)", level=1)
doc.add_paragraph("Este es un párrafo en la primera sección.")

doc.save(r'C:\Users\ANALISTAUP29\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada\Am prueba v4.docx')
