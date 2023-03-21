#!/usr/bin/env python3
# -*- coding: utf-8 -*-

#%reset -f   #Para borrar las variables

"""
*********************************************************************************************
Ayuda Memoria corta con información de presupuestal e información de otras unidades de OPEP
*********************************************************************************************
   
"""

"""
Lista de los paquetes por instalar:
"""
 
#pip install python-docx ##Paquetes para crear documentos formato "docx"
#pip install nums_from_string ##Paqute para extraer numeros de una cadena
#pip install pyjanitor ##Es una implementación de Python del paquete R janitory proporciona una API limpia para limpiar datos.

import docx
import pandas as pd
import numpy as np
import nums_from_string
import os #Este paquete permite crear actividades dependientes del sistema operativo, por ejemplo crear carpetas, conocer sobre un proceso, finalizar un proceso, etc.
import getpass
import glob
import matplotlib.pyplot as plt
from datetime import datetime
#from pyprojroot import here
import pyodbc
from janitor import clean_names # pip install pyjanitor
from pathlib import Path
from docx.shared import Pt
from docx.shared import Inches
import re #Nos porporciona opciones de coincidencia.
from docx.shared import Cm # para incluir imagenes en el documento Word
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import matplotlib.ticker as mtick
from docxtpl import DocxTemplate
from docxtpl import InlineImage
import win32com.client as win32
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

###############################################################################
# Ruta del proyecto #
###############################################################################

#Path: delmódulo pathlib, sirve para manipular rutas locales.
#getuser(): Muestra el nombre de inicio de sesión del usuario. Esta función verifica las variables de entorno LOGNAME, USER, LNAME y USERNAME, en orden, y devuelve el valor de la primera cadena no vacía.
#getpass(): Se utiliza para solicitar a los usuarios que utilicen la solicitud de cadena y lee la entrada del usuario como Contraseña. La lectura de entrada predeterminada es "Contraseña:" y se devuelve a la persona que llama como una cadena

# A continaución se comparan los "Usuarios" de la PC en donde se correra el SCRIPT, se carga la ruta segun el usuario que inicia sesion en la PC

"""
if getpass.getuser() == "analistaup29": # PC Analista UP 18 Minedu
   github = Path("C:/Users/ANALISTAUP29/Documents/GitHub/AM-python-docx")
   proyecto = Path("B:/OneDrive - Ministerio de Educación/unidad_B/2021/4. Herramientas de Seguimiento/13.AM_automatizada")
elif getpass.getuser() == "VLADIMIR": # PC Analista UP 18 Minedu
     github = Path("C:/Users/VLADIMIR/Documents/GitHub/AM-python-docx")
     proyecto = Path("C:/Users/VLADIMIR/Documents/GitHub/AM-python-docx")                    
elif getpass.getuser() == "bran": # PC Brandon
     github = Path("/Users/bran/GitHub/AM-python-docx")
     proyecto = Path("/Users/bran/GitHub/AM-python-docx")
elif getpass.getuser() == "llan_": # PC Llan - Casa 
     github = Path("D:/MINEDU_2022/GITHUB/AM-python-docx")
     proyecto = Path("D:/MINEDU_2022/GITHUB/AM-python-docx")
"""

user = getpass.getuser() # Capturo el usuario local / Ejemplo "Llan_"
user.upper()

if   user== "ANALISTAUP29": # PC Analista UP 29 Minedu
     #github = Path(r"C:\Users\ANALISTAUP29\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada")
     proyecto = Path(r"C:\Users\ANALISTAUP29\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada")
     #path_grafico= proyecto + "\graficos"  #Para guardar los graficos 
     #path_mapas= proyecto + "\mapas"  #Para guardar los graficos
     
elif user== "llan_": # PC casa Llan
     #github = Path(r"C:\Users\llan_\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada")
     proyecto = Path(r"C:\Users\llan_\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada")  
     #path_grafico= proyecto + "\graficos"  #Para guardar los graficos 
     #path_mapas= proyecto + "\mapas"  #Para guardar los graficos  #Para guardar los graficos 

elif user== "VLADIMIR": # PC Vladimir
     github = Path("C:/Users/VLADIMIR/Documents/GitHub/AM-python-docx")
     proyecto = Path("C:/Users/VLADIMIR/Documents/GitHub/AM-python-docx")   
                 
elif user == "bran": # PC Brandon
     github = Path("/Users/bran/GitHub/AM-python-docx")
     proyecto = Path("/Users/bran/GitHub/AM-python-docx")

###############################################################################
# Fechas de corte #
###############################################################################

'''
En esta sección se consigna las fechas que serviran para cargar la información y se crean variables que nos permitan
incorporar información en los textos de la AM.

Cada corte de información actualizada implica que se cambien manualmente los valores de dichas fechas en las variables
que se consignan en esta sección.
'''

# Importamos los nombres de los archivos dentro de la carpeta input
lista_archivos = os.listdir(Path(proyecto, "dataset"))

# Fecha actual
fecha_actual = datetime.today().strftime('%d-%m-%y')
nyear_actual=datetime.today().year
nmeses_actual=datetime.today().month #Rescato el número de mes que estamos hoy
ndia_actual=datetime.today().day

## A) Fecha disponibilidad
fecha_corte_disponibilidad = "20230319"
nyear_disponibilidad=fecha_corte_disponibilidad[0:4]
nmes_disponibilidad=fecha_corte_disponibilidad[4:6]
ndia_disponibilidad=fecha_corte_disponibilidad[6:]
fecha_corte_disponibilidad_format=ndia_disponibilidad + "/" + nmes_disponibilidad + "/" + nyear_disponibilidad

nmes_disp_entero=int(nmes_disponibilidad)
nmes_dispo=['enero', 'febrero', 'marzo', 'abril','mayo', 'junio','julio','agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']

## B) Fecha Nexus
fecha_corte_nexus = "20230301"
nyear_nexus=fecha_corte_nexus[0:4]
nmes_nexus=fecha_corte_nexus[4:6]
ndia_nexus=fecha_corte_nexus[6:]
fecha_corte_nexus_format=ndia_nexus + "/" + nmes_nexus + "/" + nyear_nexus

nmes_nex_entero=int(nmes_nexus)
nmes_nex=['enero', 'febrero', 'marzo', 'abril','mayo', 'junio','julio','agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']

## C) Compromisos de desempeño
fecha_corte_compromisos = "20211128"

###############################################################################
# Creación de carpeta donde se guardan los outputs #
###############################################################################

'''
Se crea la carpeta en dondes se almacenará la información correspondiente a las regiones y la fecha de corte.
'''

# Creación de carpeta
dir = os.path.join(proyecto, f"output/AM_{fecha_actual}")  #Se unen las rutas, f": Lo usamos para concatenar cadenas y valores

if not os.path.exists(dir):
    os.mkdir(dir)  # Para crear una carpeta de acuerdo a lo especificado por "dir"
    print("Se creó una nueva carpeta")
else:
    print("Ya existe la carpeta")
        
# Path de nueva carpeta
nueva_carpeta = Path(proyecto/ f"output/AM_{fecha_actual}")
path_grafico= Path(proyecto/ f"graficos")
path_mapas= Path(proyecto/ f"mapas")

###############################################################################
# Transformación de Datasets                                                  #
###############################################################################

'''
Cargamos la información de las regiones por pliego, para poder capturar esta variable en las demás bases que trabajemos

En esta base se cuenta con:
    region
    COD_PLIEGO
'''

# Base de datos región
## Cargamos nombres de regiones
nombre_regiones = pd.read_excel(proyecto / "dataset/nombre_regiones.xlsx")

# Base de datos codigo de pliego, ejecutora y ugel
nombre_otros = pd.read_excel(proyecto / "dataset/base_ue_ugel_ubigeo_2023_v2.xlsx")
#nombre_otros.loc[~(nombre_otros['CODOOII'].isnull())]
#nombre_otros.loc[~(((nombre_otros['PLIEGO']==457) & (nombre_otros['EJECUTORA']==301)) | ((nombre_otros['PLIEGO']==464) & (nombre_otros['EJECUTORA']==301)))]
nombre_otros=nombre_otros.drop_duplicates(subset = ['PLIEGO','EJECUTORA'])
nombre_otros = clean_names(nombre_otros) # Normalizamos nombres
nombre_otros.rename(columns={'pliego':'cod_pliego'},inplace=True)
nombre_otros.rename(columns={'ejecutora':'cod_ue'},inplace=True)

###############################
# Base de disponibilidad  #
###############################

'''
La base de disponbilidad contiene información SIAF (PIA, PIM, Devengado, etc) respecto a la ejecución de las intervenciones pedagógicas del MINEDU

El nombre de la base tiene la siguiente estructura:
    
    Disponibilidad_Presupuestal_"yyyy/mm/dd"

1. Se debera de homogenizar las variables de la BD genera y a nivel de intervenciones (las variables deberan tener el mismo nombre)
2. La variable de transferencia para intervenciones en la BD debera tener la siguiente estructura:
  
    tramo1_transferencia ...
    tramo2_transferencia ...
    *
    *
    *
'''

# A) Base de disponibilidad
## Cargamos base de disponibilidad
data_intervenciones = pd.read_stata(proyecto / f"dataset/bd_disponibilidad/Disponibilidad_Presupuestal_{fecha_corte_disponibilidad}_interv.dta")   
data_intervenciones = clean_names(data_intervenciones) # Normalizamos nombres
data_intervenciones['intervencion'] = data_intervenciones['intervencion'].apply(lambda x: x.strip()) #Para poder quitar los espaicos en blanco antes y despues de las cadenas

# Eliminamos filas de "No hay Intervenciones pedagogicas"
data_intervenciones = data_intervenciones[~(data_intervenciones['cod_intervencion'].isnull())]
# No consideramos: "COAR", "Juegos deportivos" y "Concursos escolares". Juegos y concursos por ahora no se cuenta con un costo aprobado por jefatura.
data_intervenciones = data_intervenciones[~((data_intervenciones['cod_intervencion']==12) | (data_intervenciones['cod_intervencion']==13) | 
                                          (data_intervenciones['cod_intervencion']==26))]
# Eliminamos  Vacaciones Truncas (Si esque hubiese)
data_intervenciones = data_intervenciones[data_intervenciones['corr'] != "3.2.8.1.5"]

# Variable para capturar la transferencias y costos, de tal manera que sumemos las variables que tengamos a la fecha de corte:
#-----------------------------------------------------
columns_names = data_intervenciones.columns.values  
columns_names_list = list(columns_names) #Todas las columnas de la base 
# transferencia_list = list(filter(lambda x: 'transferencia' in x, columns_names)) #Lista de variables de transferencia
# costos_list = list(filter(lambda x: 'costo' in x, columns_names)) #Lista de variables de costos
#contador=len(transferencia_list)

# 1. Transferencia:
transferencia_inter=[]
#id_t=['1','segunda','tercera','cuarta','quinta','sexta','septima','octava','novena','decima']
id_t=list(range(1,5)) #Considere una lista de 1 a 4 / pues en el año por lo general no existen mas de tres transferencias

for j in id_t:
     transferencia_inter =transferencia_inter + list(filter(lambda x: f'tramo{j}_transferencia' in x, columns_names)) 

data_intervenciones["transferencia_t"]=0 
i=0
for j in transferencia_inter:
    data_intervenciones["transferencia_t"]=data_intervenciones["transferencia_t"]+data_intervenciones[transferencia_inter[i]]
    i=i+1

# Contando la cantidad de transferencias:
n_transf=len(transferencia_inter)
n_transf_2=list(range(1,len(transferencia_inter)+2))

if n_transf==0:
    data_intervenciones['tramo1_transferencia'] = 0 
    data_intervenciones['tramo2_transferencia'] = 0 
elif n_transf==1:
    data_intervenciones['tramo2_transferencia'] = 0 
elif n_transf==2:
    data_intervenciones['tramo3_transferencia'] = 0 

#rescatar los nombres de variables
columns_names = data_intervenciones.columns.values  
columns_names_list = list(columns_names)

transferencia_inter=[]
for j in n_transf_2:
     transferencia_inter =transferencia_inter + list(filter(lambda x: f'tramo{j}_transferencia' in x, columns_names)) 

data_intervenciones["transferencia_t"]=0 
i=0
for j in transferencia_inter:
    data_intervenciones["transferencia_t"]=data_intervenciones["transferencia_t"]+data_intervenciones[transferencia_inter[i]]
    i=i+1

#creare este artificio, asumire que existen 2 transferencias:
#data_intervenciones['tramo1_transferencia'] = 2000 
#data_intervenciones['tramo2_transferencia'] = 5000 

# 2. Costos: nmeses_actual  
#var_costos=[]
#id_meses=['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre']
#for j in id_meses:
#   var_costos =var_costos + list(filter(lambda x: f'costo_{j}' in x, columns_names)) 

var_costos=['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre']
data_intervenciones["costo_corte"]=0 
i=0
for j in range(nmeses_actual-1):
    data_intervenciones["costo_corte"]=data_intervenciones["costo_corte"]+data_intervenciones[var_costos[i]]
    i=i+1

# Para colocar el mes en el texto de interverciones:
#id_mes_int=var_costos[int(nmes_disponibilidad)-1]

#---------------------------------------------------  
#####Generando las tablas para la AM
#---------------------------------------------------  

### Tabla1: Costo anual de las intervenciones y recursos disponibles en el PIA
#--------------------------------------------------------------------------
data_intervenciones_1 = data_intervenciones[["region","cod_pliego","nom_pliego","cod_ue","nom_ue",
                                                 "pia", "pim",'devengado',"comprometido_anual","enero", "febrero", "marzo", "abril", "mayo", 
                                                 "junio", "julio", "agosto", "septiembre", "octubre","noviembre", "diciembre","costo_upp_total",f"{transferencia_inter[0]}",f"{transferencia_inter[1]}","transferencia_t","costo_corte"]]. \
    groupby(by = ["region", "cod_pliego","cod_ue"] , as_index=False).sum()


data_intervenciones_1=data_intervenciones_1.merge(right=nombre_otros, how="left", on=["cod_pliego","cod_ue"] )
data_intervenciones_1=data_intervenciones_1[[ 'region','nom_ue','costo_upp_total','pia', 'pim', 'devengado', 'transferencia_t']]

#data_intervenciones_1.to_excel(r'C:\Users\ANALISTAUP29\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada\output\Base_preliminar.xlsx')

# ejecución
data_intervenciones_1['ejecucion']=data_intervenciones_1['devengado'] / data_intervenciones_1['pim']
#data_intervenciones_1.loc[data_intervenciones_1['ejecución'].isnull(),'ejecución']=0

## Reemplazamos inf por NaN: Se esta remplazando los valores infinitos de las divisiones que se realizaron
data_intervenciones_1.replace([np.inf, -np.inf], np.nan, inplace=True)
## Reemplazamos NaN por 0
data_intervenciones_1['ejecucion'] = data_intervenciones_1['ejecucion'].fillna("0").astype(float)


### Tabla2: a nivel de ejecutora
#--------------------------------------------------------------------------
data_intervenciones_2 = data_intervenciones[["region","cod_pliego","nom_pliego","cod_ue","nom_ue",
                                                 "pia", "pim",'devengado',"comprometido_anual","enero", "febrero", "marzo", "abril", "mayo", 
                                                 "junio", "julio", "agosto", "septiembre", "octubre","noviembre", "diciembre","costo_upp_total",f"{transferencia_inter[0]}",f"{transferencia_inter[1]}","transferencia_t","costo_corte"]]. \
    groupby(by = ["region", "cod_pliego","cod_ue"] , as_index=False).sum()

data_intervenciones_2=data_intervenciones_2.merge(right=nombre_otros, how="left", on=["cod_pliego","cod_ue"] )
data_intervenciones_2=data_intervenciones_2[['region','nom_ue', 'pim', f"{transferencia_inter[0]}",f"{transferencia_inter[1]}", 'transferencia_t']]


data_intervenciones_2_mat = data_intervenciones[["region","cod_pliego","nom_pliego","cod_ue","nom_ue", 'intervencion',
                                                 "pia", "pim",'devengado',"comprometido_anual","enero", "febrero", "marzo", "abril", "mayo", 
                                                 "junio", "julio", "agosto", "septiembre", "octubre","noviembre", "diciembre","costo_upp_total",f"{transferencia_inter[0]}",f"{transferencia_inter[1]}","transferencia_t","costo_corte"]]. \
    groupby(by = ["region", "cod_pliego","cod_ue", 'intervencion'] , as_index=False).sum()

data_intervenciones_2_mat=data_intervenciones_2_mat.merge(right=nombre_otros, how="left", on=["cod_pliego","cod_ue"] )
data_intervenciones_2_mat=data_intervenciones_2_mat[['region','nom_ue','intervencion','pim', f"{transferencia_inter[0]}",f"{transferencia_inter[1]}", 'transferencia_t']]
data_intervenciones_2_mat=data_intervenciones_2_mat.loc[data_intervenciones_2_mat['intervencion']=='Distribución de materiales educativos']

### Tabla a nivel de intervención y ejecución
#--------------------------------------------------------------------------
data_intervenciones_3 = data_intervenciones[["region","cod_pliego","nom_pliego","cod_ue","nom_ue", "cod_intervencion", "intervencion",
                                                 "pia", "pim",'devengado',"comprometido_anual","enero", "febrero", "marzo", "abril", "mayo", 
                                                 "junio", "julio", "agosto", "septiembre", "octubre","noviembre", "diciembre","costo_upp_total",f"{transferencia_inter[0]}",f"{transferencia_inter[1]}","transferencia_t","costo_corte"]]. \
    groupby(by = ["region", "cod_intervencion","intervencion"] , as_index=False).sum()


data_intervenciones_3=data_intervenciones_3[['region', 'cod_intervencion', 'intervencion', 'costo_upp_total','pim','devengado']]
#data_intervenciones_3['ejecucion']=data_intervenciones_3['devengado'] / data_intervenciones_3['pim']
#data_intervenciones_3.loc[data_intervenciones_3['ejecucion'].isnull(),'ejecucion']=0

## Reemplazamos inf por NaN: Se esta remplazando los valores infinitos de las divisiones que se realizaron
#data_intervenciones_3.replace([np.inf, -np.inf], np.nan, inplace=True)
## Reemplazamos NaN por 0
#data_intervenciones_3['ejecucion'] = data_intervenciones_3['ejecucion'].fillna("0").astype(float)

###############################
# Base de nexus PEAS CAS  #
###############################
data_cas_nexus = pd.read_excel(proyecto / f"dataset/bd_plazas_nexus/Anexos 1 CAS AIRSHP.xlsx", sheet_name='Anexo 1')   
data_cas_nexus = clean_names(data_cas_nexus) # Normalizamos nombres

# Tabla a nivel de Pliego y UE:
data_cas_nexus_1=data_cas_nexus[['cod_pliego','cod_ue','peas_programadas','costo_estimado_anual_de_la_contratacion_s_']].groupby(by=['cod_pliego','cod_ue'], as_index=False).sum()
data_cas=data_cas_nexus_1.merge(right=nombre_regiones,how="left", on='cod_pliego')
data_cas=data_cas.merge(right=nombre_otros,how="left", on=['cod_pliego','cod_ue'])
data_cas=data_cas.rename(columns={"costo_estimado_anual_de_la_contratacion_s_": "costo_anual"}) 

# Tabla a nivel de Pliego y UE:
#data_cas_nexus_1=data_cas_nexus[['','','','','']]

'''
# Transferencias

data_intervenciones['transferencia'] = data_intervenciones["ds_n°_092_2021_ef_transferencia_convivencia"] + \
     data_intervenciones["ds_n°_169_2021_ef_1°_transferencia_intervenciones"] + data_intervenciones["ds_n°_189_2021_ef_1°_transferencia_acompanatic"] \
    + data_intervenciones["ds_n°_209_2021_ef_2°_transferencia_intervenciones"] + data_intervenciones["ds_n°_210_2021_ef_2°_transferencia_acompanatic_previo"]
'''

'''
data_intervenciones=data_intervenciones.rename(columns={"transferencia_t": "transferencia"}) #Renombramos las varables para no cambiar lo codeado

# Id 1 era transferencia:
data_intervenciones['id_transferencia_int']=data_intervenciones['transferencia']>0
data_intervenciones.replace(to_replace={"id_transferencia_int": {True:1,False:0}},inplace=True)

# Eliminamos filas con 0 PIM
#condicion_elim = (data_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"] != 0) 
#pd.value_counts(condicion_elim) # Contamos los False y True
#data_intervenciones = data_intervenciones[condicion_elim]

# Calculamos porcentajes
## Avance PIM
data_intervenciones["avance_pim"] = data_intervenciones[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"]/data_intervenciones[f"pim_reporte_siaf_{fecha_corte_disponibilidad}"]

## Avance costo actual
'''

'''
data_intervenciones["costo_actual"] = data_intervenciones["costo_enero"] + data_intervenciones["costo_febrero"] + data_intervenciones["costo_marzo"] \
    + data_intervenciones["costo_abril"] + data_intervenciones["costo_mayo"] + data_intervenciones["costo_junio"] + data_intervenciones["costo_julio"] \
        + data_intervenciones["costo_agosto"] + data_intervenciones["costo_septiembre"] + data_intervenciones["costo_octubre"]
  
'''       

'''
data_intervenciones=data_intervenciones.rename(columns={"costo_corte": "costo_actual"})  #Renombramos las varables para no cambiar lo codeado 

data_intervenciones["avance_costo"] = data_intervenciones[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"]/data_intervenciones["costo_actual"]

# Mantenemos y ordenamos columnas
data_intervenciones = data_intervenciones[['region', "intervencion_pedagogica", "pia", "transferencia", f"pim_reporte_siaf_{fecha_corte_disponibilidad}",
                                           f"comprometido_anual_reporte_siaf_{fecha_corte_disponibilidad}", f"devengado_reporte_siaf_{fecha_corte_disponibilidad}", 
                                           "avance_pim", "costo_actual", "avance_costo", "id_transferencia_int"]]

## Reemplazamos inf por NaN: Se esta remplazando los valores infinitos de las divisiones que se realizaron
data_intervenciones.replace([np.inf, -np.inf], np.nan, inplace=True)

## Reemplazamos NaN por 0
data_intervenciones['avance_pim'] = data_intervenciones['avance_pim'].fillna("0").astype(float)
data_intervenciones['avance_costo'] = data_intervenciones['avance_costo'].fillna("0").astype(float)
'''

'''
#######################
# Siaf de mascarillas #
#######################
'''

'''
La base de mascarillas tiene el siguiente nombre.

Reporte_PowerBI_SIAF_"yyyy/mm/dd"

Es esta sección se realizara la carga y transformación de las variables para poder crear el word que se elabora mas adelante en el
script.

Se debe tener en cuenta que este año no se realizo transferencia para mascarilla, el recurso se pinto en el PIA. Para años posteriores, esto
no podria ser igual y puede merecr una actualización de esta sección.

'''
'''
## Cargamos la base insumo de mascarillas
data_mascarillas = pd.read_excel(proyecto / f"input/am/Mascarilla/Reporte_PowerBI_SIAF_{fecha_corte_mascarillas}.xlsx", sheet_name='Sheet1')
data_mascarillas = clean_names(data_mascarillas) # Normalizamos nombres

# Mantenemos variables de interés (PIA, PIM,  CERTIFICADO, COMPROMETIDO y DEVENGADO) y 
# colapsamos a nivel de Region y UE
data_mascarillas = data_mascarillas[["region","nom_ue","pia","pim","certificado","comprometido_anual","devengado"]]. \
   groupby(by = ["region", "nom_ue"], as_index=False).sum()

#data_mascarillas["region"] = data_mascarillas["region"].str.split(". ", n=1).apply(lambda l: "".join(l[1]))

#data_mascarillas["UNIDAD EJECUTORA"]=data_mascarillas["nom_ue"]
#data_mascarillas["RECURSOS TRANSF. (*)"]=data_mascarillas["transferencia"]
#data_mascarillas["PIM"]=data_mascarillas["pim"]
#data_mascarillas["CERT. (%)"]=data_mascarillas["certificado"]/data_mascarillas["pim"]
#data_mascarillas["COMPRO. (%)"]=data_mascarillas["comprometido_anual"]/data_mascarillas["pim"]
#data_mascarillas["DEVENGADO (%)"]=data_mascarillas["devengado"]/data_mascarillas["pim"]
#data_mascarillas["AVANCE TRANSFERENCIA (%)"]=data_mascarillas["devengado"]/data_mascarillas["transferencia"]

'''
'''
# Generaremos una base a nvel de región:
data_mascarillas_region = data_mascarillas[["region","transferencia","pim","certificado","comprometido_anual","devengado"]]. \
   groupby(by = ["region"], as_index=False).sum()   

data_mascarillas_region["RECURSOS TRANSF. (*)"]=data_mascarillas_region["transferencia"]
data_mascarillas_region["PIM"]=data_mascarillas_region["pim"]
data_mascarillas_region["CERT. (%)"]=data_mascarillas_region["certificado"]/data_mascarillas_region["pim"]
data_mascarillas_region["COMPRO. (%)"]=data_mascarillas_region["comprometido_anual"]/data_mascarillas_region["pim"]
data_mascarillas_region["DEVENGADO (%)"]=data_mascarillas_region["devengado"]/data_mascarillas_region["pim"]
data_mascarillas_region["AVANCE TRANSFERENCIA (%)"]=data_mascarillas_region["devengado"]/data_mascarillas_region["transferencia"]
'''

'''
# collapsamos la base a nivel de región para poder generar los valores de interes más adelante
# en la base no hay valores NaN.
data_mascarillas = data_mascarillas[["region","nom_ue","pia","pim","certificado","comprometido_anual","devengado"]]. \
   groupby(by = ["region"], as_index=False).sum()

# Para colocar el mes en el texto de interverciones:
id_mes_masca=id_meses[int(nmes_mascarilla)-1]
'''

############################
# Compromisos de desempeño #
############################

'''
En el año fiscal se realizán tres trasnferencias(tramos) de CdD, y una cuarta que esta ligado al buen desempeño.

El nombre de la base tiene la siguiente estructura:
    
    regiones_BD_CDD_"yyyy/mm/dd".xlsx
    
'''

'''
## Cargamos data de compromisos de desempeño
data_cdd = pd.read_excel(proyecto / f"input/compromisos_desempeno/regiones_BD_CDD_{fecha_corte_compromisos}.xlsx")

data_cdd = clean_names(data_cdd) # Normalizamos nombres

#data_cdd["pliego"] = data_cdd["pliego"].str.split(". ", n=1, expand=True) #No corria tal como estaba se cambio como se muestra en la siguinete linea
data_cdd["cod_pliego"] = data_cdd["pliego"].str.split(". ", n=1).apply(lambda l: "".join(l[0]))
data_cdd['cod_pliego'] = data_cdd['cod_pliego'].astype('int64') # Convertimos ubigeo a integer

'''

'''
# Corregimos genericas
data_cdd["generica"] = data_cdd["generica"].replace("3. BIENES Y SERVICIOS", "2.3. BIENES Y SERVICIOS")
data_cdd["generica"] = data_cdd["generica"].replace("6. ADQUISICION DE ACTIVOS NO FINANCIEROS", "2.6. ADQUISICION DE ACTIVOS NO FINANCIEROS")	
'''

'''
# Hacemos merge con base de datos región
data_cdd = data_cdd.merge(right = nombre_regiones, how="left", on = "cod_pliego")
'''

'''
# Mantenemos variables de interés (transferencia,  CERTIFICADO, COMPROMETIDO y DEVENGADO) y 
# colapsamos a nivel de Region y UE
data_cdd = data_cdd[["region", "unidad_ejecutora", "programa_presupuestal", "generica", "monto", "ds_085_2021_ef", "ds_218_2021_ef", "ds_220_2021_ef"]]. \
    groupby(by = ["region", "programa_presupuestal", "generica"], as_index=False).sum()

'''

'''
data_cdd = data_cdd. \
    groupby(by = ["region"], as_index=False).sum()

'''

'''
data_cdd_2021 = pd.read_excel(proyecto / "input/am/CDD/CDD_SIAF_20211128.xlsx", sheet_name="BD")
data_cdd_2021 = clean_names(data_cdd_2021) # Normalizamos nombres
   
# Capturamos los nombres de las columnas de la base de datos   
columns_names = data_cdd_2021.columns.values  
columns_names_list = list(columns_names) #Todas las columnas de la base 
transferencia_cdd_t1 = list(filter(lambda x: 'tramo_1' in x, columns_names))
transferencia_cdd_t2 = list(filter(lambda x: 'tramo_2' in x, columns_names))
transferencia_cdd_t3 = list(filter(lambda x: 'tramo_3' in x, columns_names))
transferencia_cdd_t4 = list(filter(lambda x: 'tramo_4' in x, columns_names))
transferencia_cdd_tramos=transferencia_cdd_t1+transferencia_cdd_t2+transferencia_cdd_t3+transferencia_cdd_t4

data_cdd_2021['Tranferencia_total']=0
for j in range(3):
     data_cdd_2021[f'Tranferencia_tramo_{j+1}'] = data_cdd_2021[transferencia_cdd_tramos[j]]
     data_cdd_2021['Tranferencia_total'] =data_cdd_2021['Tranferencia_total'] + data_cdd_2021[f'Tranferencia_tramo_{j+1}'] 
  
 '''
   
'''
variables:
    Tranferencia_total
    Tranferencia_tramo_1
    Tranferencia_tramo_2
    Tranferencia_tramo_3
    Tranferencia_tramo_4
'''
'''
data_cdd_2021 = data_cdd_2021.\
    groupby(by = ["region"], as_index=False).sum()
'''

######################################################
# Sobre el financiamiento de conceptos remunerativos #
######################################################
'''
# C) Base de Encargaturas
df_consolidado_enc = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'ENC-CONSOLIDADO-VF') 
df_consolidado_enc.fillna(0, inplace =  True) #inplace =  True: Permite que todo los valores NaN se remplacen por cero
df_consolidado_enc['COSTO'] = df_consolidado_enc['COSTO-TRAMO I']
df_consolidado_enc['PROGRAMADO POR MINEDU'] = df_consolidado_enc['APM2021'] + df_consolidado_enc['INCREMENTOS']+ df_consolidado_enc['NM ENCARGATURAS']
df_consolidado_enc.rename(columns={'DIPLOMA_GORE':'PROGRAMADO POR EL PLIEGO REGIONAL',
                                   'TRANSFERENCIA DS 217':'TRANSFERENCIA POR DS N° 217-2021-EF',
                                   'UNIDADEJECUTORA':'UNIDAD EJECUTORA'}, inplace=True) #,inplace=True: para que se renombre sin necesidad de crear otro data frame

df_consolidado_enc['APMeINCREMENTOS'] = df_consolidado_enc['APM2021'] + df_consolidado_enc['INCREMENTOS']

tabla_encargaturas = df_consolidado_enc[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                                         'PROGRAMADO POR EL PLIEGO REGIONAL', 'TRANSFERENCIA POR DS N° 217-2021-EF',
                                         'APMeINCREMENTOS','NM ENCARGATURAS']]

tabla_encargaturas_resumen = df_consolidado_enc[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU',
                                         'PROGRAMADO POR EL PLIEGO REGIONAL', 'TRANSFERENCIA POR DS N° 217-2021-EF',
                                         'APMeINCREMENTOS','NM ENCARGATURAS']]


# D) Base de Asignaciones Temporales
df_consolidado_at = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'AT-CONSOLIDADO-VF')   
df_consolidado_at.fillna(0, inplace =  True)      
df_consolidado_at.rename(columns={'REGIÓN':'REGION',
                                  'COSTO-TRAMO I':'COSTO',
                                  'APM':'PROGRAMADO POR MINEDU',
                                  'DIPLOMA_GORE':'PROGRAMADO POR EL PLIEGO REGIONAL',
                                  'TRANSFERENCIA DS 187':'TRANSFERENCIA POR DS N° 187-2021-EF'
    },inplace=True)

df_at=df_consolidado_at[['REGION','UNIDAD EJECUTORA','COSTO','PROGRAMADO POR MINEDU','TRANSFERENCIA POR DS N° 187-2021-EF']]

# E) Base de Beneficios Sociales
df_consolidado_bf = pd.read_excel(proyecto / 'input/conceptos_remunerativos/CONCEPTOS_CONSOLIDADOS_20211003.xlsx',sheet_name = 'BS-CONSOLIDADO-VF')           
df_consolidado_bf.fillna(0,inplace = True)
df_consolidado_bf['COSTO BENEFICIARIOS 2020 Y 2021'] = df_consolidado_bf['LISTAS-2021'] + df_consolidado_bf['TRAMO I-BS 2020'] + df_consolidado_bf['TRAMO II-BS 2021']
df_consolidado_bf.rename(columns={'REGIÓN':'REGION',
                                  'COSTO BENEFICIARIOS 2020 Y 2021':'COSTO',
                                  'APM':'PROGRAMADO POR MINEDU',
                                  'TRANSFERENCIA DS 072-BS 2020':'TRANSFERENCIA POR DS N° 072-2021-EF',
                                  'TRANSFERENCIA DS 256-BS 2021':'TRANSFERENCIA POR DS N° 256-2021-EF'
    },inplace=True)

df_bs = df_consolidado_bf[['REGION','UNIDAD EJECUTORA','COSTO',
                           'PROGRAMADO POR MINEDU',
                           'TRANSFERENCIA POR DS N° 072-2021-EF',
                           'TRANSFERENCIA POR DS N° 256-2021-EF']]

df_transferencia = pd.read_excel(proyecto / 'input/normas_transferencias/TRANSFERENCIAS 2021.xlsx',sheet_name = 'TRANSFERENCIAS')           
df_transferencia.fillna(0,inplace = True)
df_transferencia = clean_names(df_transferencia)
df_transferencia = df_transferencia[['region', 'norma_de_transferencia', 'concepto', 'monto_transferido']].\
groupby(by = ["region", 'norma_de_transferencia', 'concepto'] , as_index=False).sum()


normas = ["DECRETO DE URGENCIA N 065-2021", "DECRETO SUPREMO N 044-2021-EF", "DECRETO SUPREMO N 078-2021-EF"]


df_transferencia = df_transferencia.loc[df_transferencia['norma_de_transferencia'].isin(normas)] #Se esta evaluando en cada dato de la variable "norma_de_transferencia2, solo nos quedamos con los datos de al variable "norma"
df_transferencia["concepto"].replace({"CONTRATACIÓN MINDEF": "Contratación de plazas docentes en instituciones educativas de educación básica del Ministerio de Defensa"}, inplace=True)
'''
'''
######################################################
# Sobre el proceso de racionalización #
######################################################

data_creacion = pd.read_excel(proyecto / "input/creaciones/plazas_creacion_racio_2021.xlsx", sheet_name="BD",  skiprows = 2)
data_creacion = clean_names(data_creacion) # Normalizamos nombres
data_creacion = data_creacion[['d_region', 'd_dreugel', 'nivel', 'creacion_total']].\
groupby(by = ["d_region", 'd_dreugel', 'nivel'] , as_index=False).sum()
data_creacion['d_region'] = data_creacion['d_region'].str.split(r'DRE ').str[-1]
data_creacion.loc[data_creacion['d_region']=="DE DIOS", 'd_region'] = "MADRE DE DIOS"


data_creacion = data_creacion.rename(columns={'d_region':'region', 'd_dreugel':'ugel'})
data_creacion.loc[data_creacion['nivel']=="inicial", 'inicial'] = data_creacion['creacion_total']
data_creacion.loc[data_creacion['nivel']=="primaria", 'primaria'] = data_creacion['creacion_total']
data_creacion.loc[data_creacion['nivel']=="secundaria", 'secundaria'] = data_creacion['creacion_total']

data_creacion = data_creacion[['region', 'ugel','inicial', 'primaria', 'secundaria', 'creacion_total']].\
groupby(by = ["region", 'ugel'] , as_index=False).sum()

data_creacion.fillna(0, inplace =  True)

## Creación plazas docentes -PEM 2021

data_creacion_pem = pd.read_excel(proyecto / "input/creaciones/plazas_creacion_pem_2021.xlsx")
data_creacion_pem = clean_names(data_creacion_pem) # Normalizamos nombres
data_creacion_pem = data_creacion_pem[['d_region', 'd_dreugel', 'modalidad', 'req_doc', 'req_bolsa', 'req_director', 'req_subdir']].\
groupby(by = ["d_region", 'd_dreugel', 'modalidad'] , as_index=False).sum()
data_creacion_pem['d_region'] = data_creacion_pem['d_region'].str.split(r'DRE ').str[-1]
data_creacion_pem.loc[data_creacion_pem['d_region']=="DE DIOS", 'd_region'] = "MADRE DE DIOS"

data_creacion_pem = data_creacion_pem.rename(columns={'d_region':'region', 'd_dreugel':'ugel'})
data_creacion_pem.fillna(0, inplace =  True)

filtro_ebr = data_creacion_pem['modalidad']=="EBR"
creacion_ebr_pem = data_creacion_pem[filtro_ebr]
filtro_ebe = data_creacion_pem['modalidad']=="EBE"
creacion_ebe_pem = data_creacion_pem[filtro_ebe]

#----------------------------------------------------------------------#
## Brecha de plazas docentes
data_brecha = pd.read_excel(proyecto / "input/brecha/brecha_ugel_2020.xlsx", sheet_name="Data")
# Normalizamos nombres

# Mantenemos variables de interés
data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n', 'nom_exd_mov1', 'doc_e_c', 'brecha_net']].\
groupby(by = ["region", 'ugel'] , as_index=False).sum()
data_brecha['doc_e_n_cub_req'] = data_brecha['doc_e_n'] - data_brecha['nom_exd_mov1']
data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n_cub_req', 'nom_exd_mov1', 'doc_e_c', 'brecha_net']]
data_brecha.loc[data_brecha['brecha_net']<=0, 'req_neto'] = -1*data_brecha['brecha_net']
data_brecha.loc[data_brecha['brecha_net']>0, 'exc_neto'] = data_brecha['brecha_net']

#Cantidad de UGEL con requerimiento neto
data_brecha.loc[data_brecha['brecha_net']<0, 'cant_ugel_req'] = 1
data_brecha.loc[data_brecha['brecha_net']>0, 'cant_ugel_exc'] = 1

data_brecha_regional = data_brecha[['region', 'brecha_net', 'cant_ugel_req', 'cant_ugel_exc']].groupby(by = ["region"] , as_index=False).sum()
data_brecha_regional.loc[data_brecha_regional['brecha_net']<=0, 'brecha_net'] = -1*data_brecha_regional['brecha_net']
data_brecha_regional.loc[data_brecha_regional['brecha_net']>0, 'brecha_net'] = data_brecha_regional['brecha_net']
data_brecha_regional.fillna(0, inplace =  True)


data_brecha = data_brecha[["region", 'ugel', 'doc_req', 'doc_e', 'doc_e_n_cub_req', 'nom_exd_mov1', 'doc_e_c', 'req_neto', 'exc_neto', 'brecha_net']]
data_brecha.fillna(0, inplace =  True)


#----------------------------------------------------------------------#
## Bloqueo de plazas
data_bloqueo = pd.read_excel(proyecto / "input/bloqueo/plazas_bloqueo_2020.xlsx")
data_bloqueo = clean_names(data_bloqueo) # Normalizamos nombres
data_bloqueo.fillna(0, inplace =  True)

# Mantenemos variables de interés
data_bloqueo['cant_bloqueos'] = 1
data_bloqueo = data_bloqueo[["descreg", 'cant_bloqueos']].groupby(by = ["descreg"] , as_index=False).sum()
data_bloqueo = data_bloqueo.rename(columns={'descreg':'region'})

#----------------------------------------------------------------------#
## Deuda social
#data_deuda_social = pd.read_excel(proyecto / "input/deudas_sociales/deudas_sociales.xlsx")
#data_deuda_social = clean_names(data_deuda_social) # Normalizamos nombres
'''
'''
###############################################################################
# Nuevos cálculos
###############################################################################

#data_inversiones = pd.read_excel(proyecto / "input/am_corta/01. Transferencias inversiones.xlsx", sheet_name="Análisis2", skiprows=4)
#data_inversiones = clean_names(data_inversiones) # Normalizamos nombres

#monto_pia_2020 = "monto"
#monto_pim_2020 = "monto"
#monto_devengado_2020 = "monto"
#avance_2020 = "monto"
#monto_transferido_2020 = "monto"

#data_kit_lavamanos = pd.read_excel(proyecto / "input/am_corta/02. Kit de higiene y lavamanos.xlsx", sheet_name="Análisis", skiprows=5)
#data_kit_lavamanos = clean_names(data_kit_lavamanos) # Normalizamos nombres




#data_cdd_2021 = pd.read_excel(proyecto / "input/am_corta/03. CDD_20211109.xlsx", sheet_name="Sheet1")
#data_cdd_2021 = clean_names(data_cdd_2021) # Normalizamos nombres

#data_cdd_2021 = pd.read_excel(proyecto / "input/am_corta/02. CDD_AM Plantilla.xlsx", sheet_name="02. CDD", skiprows=5)
#data_cdd_2021 = clean_names(data_cdd_2021) # Normalizamos nombres


# Hacemos merge con base de datos región
#data_cdd_2021 = data_cdd_2021.merge(right = nombre_regiones, how="left", on = "cod_pliego")

# Mantenemos variables de interés (region,  pim, devengado y colapsamos a nivel de Region)
#data_cdd_2021 = data_cdd_2021[["region", "pim", "devengado"]]. \
#    groupby(by = ["region"], as_index=False).sum()

# Asignaciones temporales
data_asignaciones_2021 = pd.read_excel(proyecto / "input/am/08. Asignaciones Temporales.xlsx", sheet_name="Sheet2")

#Deuda social

data_deuda = pd.read_excel(proyecto / "input/am/00b_Deudas sociales.xlsx", sheet_name="Base", skiprows=1)
data_deuda = clean_names(data_deuda) # Normalizamos nombres

# Plazas
data_plazas = pd.read_excel(proyecto / "input/am/10. Plazas financiadas.xlsx", sheet_name="P_Financiadas", skiprows=3)

# Conceptos remunerativos
data_conceptos_remunerativos_2021 = pd.read_excel(proyecto / "input/am/00a_Conceptos remunerativos consolid.xlsx", sheet_name="Base agregada", skiprows=3)
data_conceptos_remunerativos_2021 = clean_names(data_conceptos_remunerativos_2021) # Normalizamos nombres

## Beneficios sociales
data_beneficios_sociales_2021 = pd.read_excel(proyecto / "input/am/09. Beneficios sociales.xlsx", sheet_name="TD_BS", skiprows=2)
data_beneficios_sociales_2021 = clean_names(data_beneficios_sociales_2021)
'''

###############################################################################
# Creación del documento en docx 
###############################################################################

# Generamos la lista de Regiones
lista_regiones = ["AMAZONAS", "ANCASH", "APURIMAC", "AREQUIPA", "AYACUCHO", "CAJAMARCA", "CUSCO", "HUANCAVELICA", "HUANUCO", "ICA", "JUNIN", "LA LIBERTAD", "LAMBAYEQUE", "LORETO", "MADRE DE DIOS", "MOQUEGUA", "PASCO", "PIURA", "PUNO", "SAN MARTIN", "TACNA", "TUMBES", "UCAYALI", "LIMA PROVINCIAS", "CALLAO"]

# For loop para cada región
for region in lista_regiones:
###############################################################################
# 1. Construcción de tablas e indicadores                                      #
###############################################################################
   
    #Año actual: nyear_actual
    year_actual = str(nyear_actual)
    '''
    ############################################
    # Tablas e indicadores beneficios sociales #
    ############################################
    
    region_seleccionada = data_beneficios_sociales_2021['region'] == region
    data_beneficios_region = data_beneficios_sociales_2021[region_seleccionada]
    
    costo_beneficios = str('{:,.0f}'.format(data_beneficios_region.iloc[0]['costo_beneficios']))
    pia_beneficios = str('{:,.0f}'.format(data_beneficios_region.iloc[0]['pia_beneficios']))
    ds_072_beneficios = str('{:,.0f}'.format(data_beneficios_region.iloc[0]['transferencia_072']))
    ds_256_beneficios = str('{:,.0f}'.format(data_beneficios_region.iloc[0]['transferencia_256']))
    ''' 
    '''
    ############################################
    # Tablas e indicadores conceptos remunerativos #
    ############################################
    
    region_seleccionada = data_conceptos_remunerativos_2021['region'] == region
    
    data_remuneracion_region = data_conceptos_remunerativos_2021[region_seleccionada]
    
    conceptos_remunerativos_2021 =  str('{:,.0f}'.format(data_remuneracion_region.iloc[0]['conceptos_remunerativos']))
    conceptos_remunerativos_2021_c =  data_remuneracion_region.iloc[0]['conceptos_remunerativos']
    '''
    '''
    ############################################
    # Tablas e indicadores asignaciones temporales #
    ############################################
    
    region_seleccionada = data_asignaciones_2021['region'] == region
    
    data_asignacion_region = data_asignaciones_2021[region_seleccionada]
    
    pia_asignacion = str('{:,.0f}'.format(data_asignacion_region.iloc[0]['pia_asignacion']))
    
    pim_asignacion = str('{:,.0f}'.format(data_asignacion_region.iloc[0]['pim_asignacion']))
    
    transferencia_asignacion = str('{:,.0f}'.format(data_asignacion_region.iloc[0]['transferencia_asignacion']))
    
    transferencia_asignacion_c = data_asignacion_region.iloc[0]['transferencia_asignacion']
    '''
    '''
    ############################################
    # Tablas e indicadores Plazas #
    ############################################
    
    region_seleccionada = data_plazas['region'] == region #Seleccionar region
    
    data_plazas_region = data_plazas[region_seleccionada]
    
    monto_plazas = str('{:,.0f}'.format(data_plazas_region.iloc[0]['monto_plazas']))
    monto_plazas_c = data_plazas_region.iloc[0]['monto_plazas']
    
    cantidad_plazas = str('{:,.0f}'.format(data_plazas_region.iloc[0]['cantidad_plazas']))
    '''
    ############################################
    # Tablas e indicadores Deuda Social #
    ############################################
    '''
    region_seleccionada = data_deuda['region'] == region #Seleccionar region
    
    data_deuda_region = data_deuda[region_seleccionada]
    
    monto_deuda_c = data_deuda_region.iloc[0]['﻿monto_deuda_social']
    monto_deuda = str('{:,.0f}'.format(data_deuda_region.iloc[0]['﻿monto_deuda_social']))
    '''
    ############################################
    # Tablas e indicadores cdd_2021 #
    ############################################
    '''
    region_seleccionada = data_cdd_2021['region'] == region #Seleccionar region
    data_transferencia_cdd_2021 = data_cdd_2021[region_seleccionada]
    
    monto_transferencia_cdd_2021_c = data_transferencia_cdd_2021.iloc[0]['Tranferencia_total'] # Valor numérico
    monto_transferencia_cdd_2021 = str('{:,.0f}'.format(data_transferencia_cdd_2021.iloc[0]['Tranferencia_total'])) # Valor con comas (string)
    cdd_ejecucion_2021 = str ('{:.1%}'.format(data_transferencia_cdd_2021["devengado"].sum()/data_transferencia_cdd_2021["pim"].sum()) if data_transferencia_cdd_2021["devengado"].sum()/data_transferencia_cdd_2021["pim"].sum()>0 else 0)
    '''
    ############################################
    # Tablas e indicadores kit/lavamanos #
    ############################################
    
    #    region_seleccionada = data_kit_lavamanos['region'] == region #Seleccionar region
    
    #    data_region_kit = data_kit_lavamanos[region_seleccionada]
    #    monto_region_kit = str('{:,.0f}'.format(data_region_kit.iloc[0]['kit_de_higiene_transferencia']))
    #    monto_region_kit_c = data_region_kit.iloc[0]['kit_de_higiene_transferencia']
    #    data_region_lavamanos = data_kit_lavamanos[region_seleccionada]
    #    monto_region_lavamanos = str('{:,.0f}'.format(data_region_lavamanos.iloc[0]['lavamanos_transferencia']))
    #    monto_region_lavamanos_c = data_region_lavamanos.iloc[0]['lavamanos_transferencia']
    #    avance_kit = data_kit_lavamanos[region_seleccionada]
    #    avance_kit = str('{:,.0f}'.format(avance_kit.iloc[0]['kit_de_higiene_declaracion']))
    #    avance_lavamanos = data_kit_lavamanos[region_seleccionada]
    #    avance_lavamanos = str('{:,.0f}'.format(avance_lavamanos.iloc[0]['lavamanos_declaracion']))
    
    
    ############################################
    # Tablas e indicadores base inversiones #
    ############################################
    
    #   region_seleccionada = data_inversiones['region'] == region #Seleccionar region
    #    data_region_inversiones = data_inversiones[region_seleccionada]
    #    monto_region_inversiones_c = data_region_inversiones.iloc[0]['monto_inversiones'] # Valor numérico
    #    monto_region_inversiones = str('{:,.0f}'.format(data_region_inversiones.iloc[0]['monto_inversiones']))
    #    fila1 = data_inversiones[region_seleccionada]
    #    fila1 = str(fila1.iloc[0]['texto_1'])
    #    fila2 = data_inversiones[region_seleccionada]
    #    fila2 = str(fila2.iloc[0]['texto_2']) 
    
    ############################################
    # Tablas e indicadores base disponibilidad #
    ############################################
    
    # Generamos los indicadores de PIM y ejecución de intervenciones
    region_seleccionada = data_intervenciones_1['region'] == region #Seleccionar region
    tabla_intervenciones = data_intervenciones_1[region_seleccionada]   
    pia_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["pia"].sum()))
    pim_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["pim"].sum()))
    ejecucion_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["devengado"].sum()))
    costo_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["costo_upp_total"].sum()))
    porcentaje_ejecucion = str('{:,.1%}'.format(tabla_intervenciones["devengado"].sum()/tabla_intervenciones["pim"].sum())  if tabla_intervenciones["devengado"].sum()/tabla_intervenciones["pim"].sum()>0 else 0)
    
    #porcentaje_costomesactual = str('{:,.1%}'.format(tabla_intervenciones[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones["costo_actual"].sum()))
    #transferencia_region_2021 = str('{:,.0f}'.format(tabla_intervenciones["transferencia"].sum()))
    #transferencia_region_2021_c = tabla_intervenciones["transferencia"].sum()
    
    # TOTAL
    tabla_intervenciones_grafico=data_intervenciones_1[region_seleccionada]
    tabla_intervenciones_formato = data_intervenciones_1[region_seleccionada]
    
    #tabla_intervenciones_formato=tabla_intervenciones_formato[[ 'region','nom_ue','costo_upp_total','pia', 'pim', 'devengado', 'transferencia_t']]
    
    # Generamos porcentaje de avance
    porcentaje_ejecucion_a = tabla_intervenciones_formato["devengado"].sum()/tabla_intervenciones_formato["pim"].sum()
    #porcentaje_costomesactual_a = tabla_intervenciones_formato[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato["costo_actual"].sum()
    
    # Generamos fila total
    total_int = tabla_intervenciones_formato.groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_intervenciones_formato = tabla_intervenciones_formato.append(total_int, ignore_index=True)
    
    # Reemplazamos % de avance pim correctos en fila total
    tabla_intervenciones_formato.iloc[-1, tabla_intervenciones_formato.columns.get_loc('ejecucion')] = porcentaje_ejecucion_a
    
    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_intervenciones_formato['nom_ue'] = tabla_intervenciones_formato['nom_ue'].fillna("Total")
    tabla_intervenciones_formato['ejecucion'] = tabla_intervenciones_formato['ejecucion'].fillna("0").astype(float)
    
    # Formato para la tabla
    formato_tabla_intervenciones = {
    "nom_ue" : "{}",
    "costo_upp_total" : "{:,.0f}",
    "pia": "{:,.0f}"
    #    "avance_pim": "{:,.1%}",
    #    "costo_actual": "{:,.0f}",
    #    "ejecucion": "{:,.1%}",
    }
    
    tabla_intervenciones_formato = tabla_intervenciones_formato.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})

    ############################################
    # Tablas e indicadores de transferencia #
    ############################################
    # Generamos los indicadores de PIM y ejecución de intervenciones
    region_seleccionada = data_intervenciones_2['region'] == region #Seleccionar region
    tabla_intervenciones = data_intervenciones_2[region_seleccionada] 
    
    #Para poder hacer mi indicador si existe transferencia
    tr1_intervenciones_num = tabla_intervenciones["tramo1_transferencia"].sum()
    tr2_intervenciones_num = tabla_intervenciones["tramo2_transferencia"].sum()
    tr_intervenciones_num = tabla_intervenciones["transferencia_t"].sum()
    
    #pia_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["pia"].sum()))
    pim_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["pim"].sum()))
    
    if tr1_intervenciones_num>0:
        tr1_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["tramo1_transferencia"].sum()))
    #else: 
    #    tr1_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["tramo1_transferencia"].sum()))
        
    if tr2_intervenciones_num>0:   
        tr2_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["tramo2_transferencia"].sum()))
    
    tr_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["transferencia_t"].sum()))
    
    #Para materiales
    region_seleccionada_2 = data_intervenciones_2_mat['region'] == region #Seleccionar region
    tabla_intervenciones_2 = data_intervenciones_2_mat[region_seleccionada_2] 
    
    #transf_materiales=tabla_intervenciones_2.loc[tabla_intervenciones_2['intervencion']=='Distribución de materiales educativos']
    transf_materiales = str('{:,.0f}'.format(tabla_intervenciones_2["transferencia_t"].sum()))
    
    #ejecucion_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["devengado"].sum()))
    #costo_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["costo_upp_total"].sum()))
    #porcentaje_ejecucion = str('{:,.1%}'.format(tabla_intervenciones["devengado"].sum()/tabla_intervenciones["pim"].sum()))
    #porcentaje_costomesactual = str('{:,.1%}'.format(tabla_intervenciones[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones["costo_actual"].sum()))
    #transferencia_region_2021 = str('{:,.0f}'.format(tabla_intervenciones["transferencia"].sum()))
    #transferencia_region_2021_c = tabla_intervenciones["transferencia"].sum()
    
    # TOTAL
    tabla_intervenciones_formato_2 = data_intervenciones_2[region_seleccionada]
    
    #tabla_intervenciones_formato=tabla_intervenciones_formato[[ 'region','nom_ue','costo_upp_total','pia', 'pim', 'devengado', 'transferencia_t']]
    
    # Generamos porcentaje de avance
    #porcentaje_ejecucion_a = tabla_intervenciones_formato["devengado"].sum()/tabla_intervenciones_formato["pim"].sum()
    #porcentaje_costomesactual_a = tabla_intervenciones_formato[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato["costo_actual"].sum()
    
    # Generamos fila total
    total_int = tabla_intervenciones_formato_2.groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_intervenciones_formato_2 = tabla_intervenciones_formato_2.append(total_int, ignore_index=True)
    
    # Reemplazamos % de avance pim correctos en fila total
    #tabla_intervenciones_formato.iloc[-1, tabla_intervenciones_formato.columns.get_loc('ejecucion')] = porcentaje_ejecucion_a
    
    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_intervenciones_formato_2['nom_ue'] = tabla_intervenciones_formato_2['nom_ue'].fillna("Total")
    #tabla_intervenciones_formato['ejecucion'] = tabla_intervenciones_formato['ejecucion'].fillna("0").astype(float)
    
    if ((tr1_intervenciones_num>0) & (tr2_intervenciones_num>0)): 
        id_trans="1" #Existen 2 transferencias
    elif ((tr1_intervenciones_num>=0) & (tr2_intervenciones_num==0)):
        id_trans="2" #Existe transferencia 1
    else: 
        id_trans="3" #No existen transferencias

    if id_trans=="1":
        # Formato para la tabla
        formato_tabla_intervenciones = {
        "nom_ue" : "{}",
        "pim" : "{:,.0f}",
        "tramo1_transferencia": "{:,.0f}",
        "tramo2_transferencia": "{:,.0f}",
        "transferencia_t": "{:,.0f}"
        #    "avance_pim": "{:,.1%}",
        #    "costo_actual": "{:,.0f}",
        #    "ejecucion": "{:,.1%}",
        }
        
        tabla_intervenciones_formato_2 = tabla_intervenciones_formato_2.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})
    
    elif id_trans=="2":
        # Formato para la tabla
        formato_tabla_intervenciones = {
        "nom_ue" : "{}",
        "pim" : "{:,.0f}",
        "tramo1_transferencia": "{:,.0f}",
        "transferencia_t": "{:,.0f}"
        #    "avance_pim": "{:,.1%}",
        #    "costo_actual": "{:,.0f}",
        #    "ejecucion": "{:,.1%}",
        }
        
        tabla_intervenciones_formato_2 = tabla_intervenciones_formato_2.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})
        

    ############################################
    # Tablas recursos asignados a nivel de las intervenciones y acciones pedagogicas  #
    ############################################
    
    # Generamos los indicadores de PIM y ejecución de intervenciones
    region_seleccionada = data_intervenciones_3['region'] == region #Seleccionar region
    tabla_intervenciones = data_intervenciones_3[region_seleccionada]   
    #pia_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["pia"].sum()))
    pim_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["pim"].sum()))
    #ejecucion_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["devengado"].sum()))
    costo_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["costo_upp_total"].sum()))
    #porcentaje_ejecucion = str('{:,.1%}'.format(tabla_intervenciones["devengado"].sum()/tabla_intervenciones["pim"].sum())  if tabla_intervenciones["devengado"].sum()/tabla_intervenciones["pim"].sum()>0 else 0)
    
    #porcentaje_costomesactual = str('{:,.1%}'.format(tabla_intervenciones[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones["costo_actual"].sum()))
    #transferencia_region_2021 = str('{:,.0f}'.format(tabla_intervenciones["transferencia"].sum()))
    #transferencia_region_2021_c = tabla_intervenciones["transferencia"].sum()
    
    # TOTAL
    #tabla_intervenciones_grafico=data_intervenciones_1[region_seleccionada]
    tabla_intervenciones_formato_3 = data_intervenciones_3[region_seleccionada]
    tabla_intervenciones_formato_4 = data_intervenciones_3[region_seleccionada]
    
    #tabla_intervenciones_formato=tabla_intervenciones_formato[[ 'region','nom_ue','costo_upp_total','pia', 'pim', 'devengado', 'transferencia_t']]
    
    #Generare un id para identificar algunas datos:
    total_intervencion = data_intervenciones_3[region_seleccionada]
    total_intervencion=total_intervencion.drop_duplicates(['cod_intervencion'])
    total_intervencion['id_int']=0
    total_intervencion.loc[total_intervencion['cod_intervencion']>0,'id_int']=1
    total_intervencion= str('{:,.0f}'.format(total_intervencion["id_int"].sum()))
    
    total_interven_mayor= data_intervenciones_3[region_seleccionada]
    total_interven_mayor = total_interven_mayor[['pim','intervencion']].groupby(by = ["intervencion"], as_index=False).sum()
    total_interven_mayor=total_interven_mayor.sort_values('pim', ascending=False)
    
    #Cinco mayores:
    total_interven_mayor_1=str('{:,.2f}'.format(round(total_interven_mayor.iloc[0,1]/1000000,2)))
    total_interven_mayor_2=str('{:,.2f}'.format(round(total_interven_mayor.iloc[1,1]/1000000,2)))
    total_interven_mayor_3=str('{:,.2f}'.format(round(total_interven_mayor.iloc[2,1]/1000000,2)))
    total_interven_mayor_4=str('{:,.2f}'.format(round(total_interven_mayor.iloc[3,1]/1000000,2)))
    total_interven_mayor_5=str('{:,.2f}'.format(round(total_interven_mayor.iloc[4,1]/1000000,2)))  
        
    # Generamos fila total
    total_int = tabla_intervenciones_formato_3.groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_intervenciones_formato_3 = tabla_intervenciones_formato_3.append(total_int, ignore_index=True)
    
    #Para cuadro de ejecucion
    tabla_intervenciones_formato_4 = tabla_intervenciones_formato_4[['region','intervencion','devengado', 'pim']].groupby(by = ['region',"intervencion"], as_index=False).sum()
    
    # Generamos porcentaje de avance
    tabla_intervenciones_formato_4['ejecucion'] = tabla_intervenciones_formato_4["devengado"]/tabla_intervenciones_formato_4["pim"]
    tabla_intervenciones_formato_4.loc[tabla_intervenciones_formato_4['ejecucion'].isnull(),'ejecucion']=0
    tabla_intervenciones_formato_4=tabla_intervenciones_formato_4.sort_values('ejecucion', ascending=False)

    # Generamos fila total
    total_int_2 = tabla_intervenciones_formato_4.groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_intervenciones_formato_4 = tabla_intervenciones_formato_4.append(total_int_2, ignore_index=True)
    tabla_intervenciones_formato_4['intervencion'] = tabla_intervenciones_formato_4['intervencion'].fillna("Total")
    tabla_intervenciones_formato_4['ejecucion'] = tabla_intervenciones_formato_4['ejecucion'].fillna("0").astype(float)
    #tabla_intervenciones_formato_3.iloc[-1, tabla_intervenciones_formato_3.columns.get_loc('ejecucion')] = porcentaje_ejecucion_a
    
    tabla_intervenciones_formato_5=tabla_intervenciones_formato_4.copy()
    tabla_intervenciones_formato_5=tabla_intervenciones_formato_5.loc[tabla_intervenciones_formato_5['intervencion']!="Total"]
    tabla_intervenciones_formato_5=tabla_intervenciones_formato_5.loc[tabla_intervenciones_formato_5['ejecucion']>0]
    lista_interve_mayor=tabla_intervenciones_formato_5.to_numpy().transpose().tolist()
    lista_interve_mayor=lista_interve_mayor[1]  #Intervenciones con devengado mayor a ce
    nombres_inte=", ".join(lista_interve_mayor)

    #tabla_intervenciones_formato_4['intervencion'] = tabla_intervenciones_formato_4['intervencion'].fillna("Total")
    #tabla_intervenciones_formato_4['ejecucion'] = tabla_intervenciones_formato_4['ejecucion'].fillna("0").astype(float)
    #tabla_intervenciones_formato_3.iloc[-1, tabla_intervenciones_formato_3.columns.get_loc('ejecucion')] = porcentaje_ejecucion_a
    
    #Incluimos palabra "total" y "-" en vez de NaN
    #para los devengados cero
    tabla_intervenciones_formato_6=tabla_intervenciones_formato_4.copy()
    tabla_intervenciones_formato_6=tabla_intervenciones_formato_6.loc[tabla_intervenciones_formato_6['ejecucion']==0]
    tabla_intervenciones_formato_6=tabla_intervenciones_formato_6.loc[tabla_intervenciones_formato_6['intervencion']!="Total"]
    lista_interve_cero=tabla_intervenciones_formato_6.to_numpy().transpose().tolist()
    lista_interve_cero=lista_interve_cero[1]  #Intervenciones con devengado mayor a cero
    nombres_inte_cero=", ".join(lista_interve_cero)

    tabla_intervenciones_formato_3['intervencion'] = tabla_intervenciones_formato_3['intervencion'].fillna("Total")
    #tabla_intervenciones_formato['nom_ue'] = tabla_intervenciones_formato['nom_ue'].fillna("Total")
    #tabla_intervenciones_formato['ejecucion'] = tabla_intervenciones_formato['ejecucion'].fillna("0").astype(float)
    tabla_intervenciones_formato_3=tabla_intervenciones_formato_3[['region','intervencion','costo_upp_total','pim']]
    
    # Formato para la tabla
    formato_tabla_intervenciones = {
    "intervencion" : "{}",
    "costo_upp_total" : "{:,.0f}",
    "pim": "{:,.0f}"
    #    "avance_pim": "{:,.1%}",
    #    "costo_actual": "{:,.0f}",
    #    "ejecucion": "{:,.1%}",
    }
    
    tabla_intervenciones_formato_3 = tabla_intervenciones_formato_3.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})

    #tabla_intervenciones_formato['nom_ue'] = tabla_intervenciones_formato['nom_ue'].fillna("Total")
    #tabla_intervenciones_formato['ejecucion'] = tabla_intervenciones_formato['ejecucion'].fillna("0").astype(float)
    tabla_intervenciones_formato_4=tabla_intervenciones_formato_4[['region','intervencion','pim','devengado','ejecucion']]
    
    # Formato para la tabla
    formato_tabla_intervenciones = {
    "intervencion" : "{}",
    "pim" : "{:,.0f}",
    "devengado": "{:,.0f}",
    "ejecucion": "{:,.1%}"
    #    "costo_actual": "{:,.0f}",
    #    "ejecucion": "{:,.1%}",
    }
    
    tabla_intervenciones_formato_4 = tabla_intervenciones_formato_4.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})
    
    ############################################
    # Tablas para Contratacion CAS Personal
    ############################################
    # Generamos los indicadores de PIM y ejecución de intervenciones
    region_seleccionada = data_cas['region'] == region #Seleccionar region
    tabla_intervenciones = data_cas[region_seleccionada]   
    total_peas_region = str('{:,.0f}'.format(tabla_intervenciones["peas_programadas"].sum()))
    cant_ue_region = str('{:,.0f}'.format(tabla_intervenciones["cod_ue"].count()))
    costo_anual_region = str('{:,.2f}'.format(tabla_intervenciones["costo_anual"].sum()/1000000))
    #costo_intervenciones_region = str('{:,.0f}'.format(tabla_intervenciones["costo_upp_total"].sum()))
    #porcentaje_ejecucion = str('{:,.1%}'.format(tabla_intervenciones["devengado"].sum()/tabla_intervenciones["pim"].sum())  if tabla_intervenciones["devengado"].sum()/tabla_intervenciones["pim"].sum()>0 else 0)
    
    #porcentaje_costomesactual = str('{:,.1%}'.format(tabla_intervenciones[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones["costo_actual"].sum()))
    #transferencia_region_2021 = str('{:,.0f}'.format(tabla_intervenciones["transferencia"].sum()))
    #transferencia_region_2021_c = tabla_intervenciones["transferencia"].sum()
    
    # TOTAL
    tabla_intervenciones_formato_6 = data_cas[region_seleccionada]
    
    #tabla_intervenciones_formato=tabla_intervenciones_formato[[ 'region','nom_ue','costo_upp_total','pia', 'pim', 'devengado', 'transferencia_t']]
    
    # Generamos porcentaje de avance
    #porcentaje_ejecucion_a = tabla_intervenciones_formato["devengado"].sum()/tabla_intervenciones_formato["pim"].sum()
    #porcentaje_costomesactual_a = tabla_intervenciones_formato[f"devengado_reporte_siaf_{fecha_corte_disponibilidad}"].sum()/tabla_intervenciones_formato["costo_actual"].sum()
    
    # Generamos fila total
    total_int = tabla_intervenciones_formato_6.groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_intervenciones_formato_6 = tabla_intervenciones_formato_6.append(total_int, ignore_index=True)
    
    # Reemplazamos % de avance pim correctos en fila total
    #tabla_intervenciones_formato_6.iloc[-1, tabla_intervenciones_formato_6.columns.get_loc('ejecucion')] = porcentaje_ejecucion_a
    tabla_intervenciones_formato_6=tabla_intervenciones_formato_6[['region','nom_ue','peas_programadas','costo_anual']]
    
    #Incluimos palabra "total" y "-" en vez de NaN
    tabla_intervenciones_formato_6['nom_ue'] = tabla_intervenciones_formato_6['nom_ue'].fillna("Total")
    #tabla_intervenciones_formato['ejecucion'] = tabla_intervenciones_formato['ejecucion'].fillna("0").astype(float)
    
    # Formato para la tabla
    formato_tabla_intervenciones = {
    "nom_ue" : "{}",
    "peas_programadas" : "{:,.0f}",
    "costo_anual": "{:,.0f}"
    #    "avance_pim": "{:,.1%}",
    #    "costo_actual": "{:,.0f}",
    #    "ejecucion": "{:,.1%}",
    }
    
    tabla_intervenciones_formato_6 = tabla_intervenciones_formato_6.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})
    
    '''
    # Generamos el ID para la transferencia
    id_transferencia_intv = str('{:,.0f}'.format(tabla_intervenciones["id_transferencia_int"].sum()))
    id_n_transf= str(n_transf)
    '''
    ############################################
    # Tablas e indicadores mascarillas #
    ############################################
    '''
    # Generamos la tabla "tabla1_mascarilla" - mantiene la región i de la lista de
    # regiones
    region_seleccionada = data_mascarillas['region'] == region
    tabla_mascarillas = data_mascarillas[region_seleccionada]
    # Generamos los indicadores de PIA, PIM, certificado, comprometido y ejecución de mascarillas
    pia_mascarilla = str('{:,.0f}'.format(tabla_mascarillas["pia"].sum()))   
    pim_mascarilla = str('{:,.0f}'.format(tabla_mascarillas["pim"].sum()))
    cert_mascarilla = str ('{:.1%}'.format(tabla_mascarillas["certificado"].sum()/tabla_mascarillas["pim"].sum()) if tabla_mascarillas["certificado"].sum()/tabla_mascarillas["pim"].sum()>0 else 0)
    comp_mascarilla = str ('{:.1%}'.format(tabla_mascarillas["comprometido_anual"].sum()/tabla_mascarillas["pim"].sum()) if tabla_mascarillas["comprometido_anual"].sum()/tabla_mascarillas["pim"].sum()>0 else 0)
    dev_mascarilla = str ('{:.1%}'.format(tabla_mascarillas["devengado"].sum()/tabla_mascarillas["pim"].sum()) if tabla_mascarillas["devengado"].sum()/tabla_mascarillas["pim"].sum()>0 else 0)
        
    #transferencia_mascarilla_c = tabla_mascarillas["transferencia"].sum()
    #transferencia_mascarilla_millones = str('{:,.1f}'.format(tabla_mascarillas["transferencia"].sum()/1000000))
    #devengado_mascarillas=str('{:.1%}'.format(tabla_mascarillas["devengado"].sum()/tabla_mascarillas["transferencia"].sum()))
    '''
    #################################################
    # Tablas e indicadores compromisos de desempeño #
    #################################################
    '''
    # Generamos la tabla "tabla_cdd" - mantiene la región i de la lista de
    # regiones
    region_seleccionada = data_cdd['region'] == region
    tabla_cdd = data_cdd[region_seleccionada]
    
    # Generamos fila total
    total_cdd = tabla_cdd[["region", "programa_presupuestal", "generica", "monto", "ds_085_2021_ef", "ds_218_2021_ef", "ds_220_2021_ef"]]. \
    groupby(by = ["region"], as_index=False).sum()
    
    # Realizamos append del total en la tabla
    tabla_cdd = tabla_cdd.append(total_cdd, ignore_index=True)
    ''' 
    ########################
    # Tablas Encargaturas #
    ########################
    '''
    region_seleccionada = df_consolidado_enc['REGION'] == region
    tabla1 = tabla_encargaturas[region_seleccionada]
    costo_enc = str('{:,.0f}'.format(tabla1["COSTO"].sum())) 
    apmeincre = str('{:,.0f}'.format(tabla1["APMeINCREMENTOS"].sum()))
    prog_gore = str('{:,.0f}'.format(tabla1["PROGRAMADO POR EL PLIEGO REGIONAL"].sum()))
    nm_enca = str('{:,.0f}'.format(tabla1['NM ENCARGATURAS'].sum()))
    ds_217 = str('{:,.0f}'.format(tabla1['TRANSFERENCIA POR DS N° 217-2021-EF'].sum()))
    ds_217_c = tabla1['TRANSFERENCIA POR DS N° 217-2021-EF'].sum()
    '''
    ###################################
    #  Tablas Asignaciones Temporales #
    ###################################
    '''
    region_seleccionada = df_at['REGION'] == region
    tabla2 = df_at[region_seleccionada]
    costo_at = str('{:,.0f}'.format(tabla2["COSTO"].sum()))
    apm_at = str('{:,.0f}'.format(tabla2["PROGRAMADO POR MINEDU"].sum())) 
    ds_187_at = str('{:,.0f}'.format(tabla2["TRANSFERENCIA POR DS N° 187-2021-EF"].sum())) 
    ds_187_at_c = tabla2["TRANSFERENCIA POR DS N° 187-2021-EF"].sum()    
    '''
    ###################################
    #    Tablas Beneficios Sociales   #
    ###################################
    '''
    region_seleccionada = df_bs['REGION'] == region
    region_seleccionada2 = df_consolidado_bf['REGION'] == region
    tabla3 = df_bs[region_seleccionada]
    tabla3_2 = df_consolidado_bf[region_seleccionada]
    costo_bs = str('{:,.0f}'.format(tabla3["COSTO"].sum()))
    apm_bs = str('{:,.0f}'.format(tabla3["PROGRAMADO POR MINEDU"].sum()))    
    ds_72_bs = str('{:,.0f}'.format(tabla3["TRANSFERENCIA POR DS N° 072-2021-EF"].sum())) 
    ds_72_bs_c = tabla3["TRANSFERENCIA POR DS N° 072-2021-EF"].sum()
    ds_256_bs = str('{:,.0f}'.format(tabla3["TRANSFERENCIA POR DS N° 256-2021-EF"].sum()))
    '''
    ########################################
    #    Tablas Financiamiento de Plazas   #
    ########################################
    '''
    region_seleccionada = data_creacion['region'] == region #Seleccionar region
    tabla_creacion = data_creacion[region_seleccionada]
    creacion_region = str('{:,.0f}'.format(tabla_creacion["creacion_total"].sum()))
    '''
    ###################################
    #    Suma de variables   #
    ###################################

    # transferencia_region_2021_c
    # ds_217_c
    # ds_187_at_c
    # ds_72_bs_c
    # monto_region_inversiones_c
    # transferencia_mascarilla_c
    # monto_region_kit_c
    # monto_region_lavamanos_c
    # monto_transferencia_cdd_2021_c
    # monto_plazas
    # monto_deuda_c
    
    #total_transferido = str('{:,.0f}'.format(np.sum([transferencia_region_2021_c, conceptos_remunerativos_2021_c, monto_deuda_c, monto_region_inversiones_c, transferencia_mascarilla_c, monto_region_kit_c, monto_region_lavamanos_c, monto_transferencia_cdd_2021_c])))
    #total_transferido = str('{:,.0f}'.format(np.sum([transferencia_region_2021_c, conceptos_remunerativos_2021_c, monto_deuda_c, transferencia_mascarilla_c, monto_transferencia_cdd_2021_c])))

        
    ############################################
    # Tablas de Carátula #
    ############################################
    
    # Conceptos remunerativos: 
        # Pago de encargaturas - ds_217
        # Pago de asignaciones temporales - ds_187_at - transferencia_asignacion_c
        # Pago de beneficios sociales - ds_72_bs
        # Financiamiento de plazas - monto_plazas_c
    
    '''
    # Creamos tabla con fechas de corte
    tabla_1 = (
        ("Intervenciones pedagógicas", transferencia_region_2021),
        ("Conceptos remunerativos", conceptos_remunerativos_2021),
        ("Deuda social", monto_deuda)
        #("Inversiones", monto_region_inversiones)
    )
    '''
    
    '''
    tabla_2 = (
        ("Mascarillas y protectores faciales", transferencia_mascarilla),
        #("Kit de higiene", monto_region_kit),
        #("Estaciones de lavado de manos", monto_region_lavamanos)
    )
    '''
    
    '''
    tabla_3 = (
        ("Compromisos de desempeño", monto_transferencia_cdd_2021),
    )
    '''
    
    '''
    tabla_4 = (
        ("TOTAL TRANSFERIDO", total_transferido),
    )
    '''
    
###############################################################################
# 2. Gráficos para incluir en el documento                                    #
###############################################################################
    
    ##################################################################################
    # Gráfico 1: Avance en la ejecución presupuestal por Unidad Ejecutora (Dev/PIM%)
    ##################################################################################
    df_collaps=tabla_intervenciones_grafico[['pim','devengado','nom_ue']].groupby(by="nom_ue", as_index=False).sum()
    df_collaps['ejecucion'] = (df_collaps["devengado"]/df_collaps["pim"])*100
    df_collaps.replace([np.inf, -np.inf], np.nan, inplace=True)
    df_collaps.loc[df_collaps["ejecucion"].isnull(), 'ejecucion']=0
        
    df_sorted = df_collaps.sort_values('ejecucion', ascending=False)
    #plt.barh(df_sorted["nom_ue"], df_sorted["ejecucion"],color=['#C00000'])
    
    fig, ax = plt.subplots()
    ax.barh(df_sorted['nom_ue'], df_sorted['ejecucion'], color=['#C00000'])
        
    # Configurar el formato del eje x
    fmt = '%.0f%%'
    xticks = mtick.FormatStrFormatter(fmt)
    plt.gca().xaxis.set_major_formatter(xticks)
    
    for i, v in enumerate(df_sorted['ejecucion']):
        plt.text(v + 0.1, i - 0.1, '{:.1f}%'.format(v), color='black', fontsize=9, fontweight='normal', fontstyle='normal', fontfamily='Arial(Cuerpo)')
    
    #plt.axvline(x=31357871, color='blue', linestyle='--', linewidth=1)
    #plt.axvline(x=31367871, color='green', linestyle='--', linewidth=1)
    
    plt.grid(axis='x', color='gray', linestyle=':', linewidth=1, zorder=-10)
    #plt.tight_layout()
    
    for spine in ax.spines.values():
        spine.set_visible(False)
  
    plt.savefig(path_grafico / "Gráfico_1.png", bbox_inches="tight")
    #plt.savefig("Gráfico_1.png") # guardamos el gráfico para cargarlo en el word
    plt.show()
    
###############################################################################
# 3. Inclusión del texto del documento                                        #
###############################################################################
    document = Document(proyecto / "formato/FORMATO_FINAL_UPP.docx") # Creación del documento en base al template
    #hoja = document.sections[2].start_page

    #document.add_page_break()
    #title = document.sections[0].header.paragraphs[0].runs[0]
    #title.add_break(docx.enum.text.WD_BREAK.LINE)
    #seccion = document.sections[2]
    #num_pagina = seccion.start_page
           
    #document.add_page_break()
    #document.add_page_break()
    #title = document.sections[0]
    #title=document.add_heading('AYUDA MEMORIA') #Título del documento
    #run = title.add_run()
    #title.add_run(' DE LA REGIÓN ')
    #title.add_run(region)
    #run = title.add_run()
    
    #for i in range(len(document.sections)):
    #    if document.sections[i].start_page == hoja:
    #        document.sections[i].header_distance = document.sections[i].header_distance + docx.shared.Cm(1.0)
    #       document.sections[i].footer_distance = document.sections[i].footer_distance + docx.shared.Cm(1.0)
    
    '''
    seccion = document.sections[2]
    encabezado = seccion.header
    texto_encabezado = "AYUDA MEMORIA"
    parrafo_encabezado = encabezado.paragraphs[0]
    parrafo_encabezado.text = texto_encabezado
    '''
    ###########################################################################
    # Carátula del documento #
    ###########################################################################

    #document.add_picture(f"/Users/bran/GitHub/AM-python-docx/input/maps/{region}.PNG", width=Inches(2.7))
    
    # --- add a 2 x 2 table as an example ---
   # table = document.add_table(rows=2, cols=2)
    # --- get the first cell of the first row ---
    #cell = table.rows[0].cells[0]
    # --- by default a cell has one paragraph with zero runs ---
    #paragraph = cell.paragraphs[0]
    # --- add a run in which to place the picture ---
    #run2 = paragraph.add_run()
    # --- add the picture to that run ---
    #run2.add_picture(f"/Users/bran/GitHub/AM-python-docx/input/maps/{region}.PNG", width=Inches(2.7))
        
    #############################################################################################################################################################################
    # Incluimos sección 1 de intervenciones pedagógicas: Tabla N° 01. Costo anual de las intervenciones y recursos disponibles en el PIA, por Unidad Ejecutora                  #
    #############################################################################################################################################################################
    document.add_heading("1. Intervenciones y Acciones Pedagógicas (IAP)", level=1) # 1) Intervenciones pedagógicas  
    #document.add_heading(f"Corte: {fecha_corte_disponibilidad_format}", level=3)
        
    interv_parrafo2 = document.add_paragraph("Mediante el numeral 42.1 del artículo 42 de la Ley N° 31638, Ley de Presupuesto del Sector Público del Año Fiscal 2023, \
se autoriza al Ministerio de Educación, a realizar modificaciones presupuestales a favor de los gobiernos regionales, \
hasta por el monto de S/ 264 531 490,00 (DOSCIENTOS SESENTA Y CUATRO MILLONES QUINIENTOS TREINTA Y UN MIL CUATROCIENTOS NOVENTA Y 00/100 SOLES),\
 para el financiamiento y ejecución de las intervenciones y acciones pedagógicas señaladas en el artículo 42 de la citada Ley.", style='List Bullet')
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY        
    #interv_parrafo2.style = document.styles['Heading 5']
    #interv_parrafo2.style = document.styles['List Bullet']
    #interv_parrafo2 = document.add_paragraph()
    
    interv_parrafo2 = document.add_paragraph(f"El Gobierno Regional de {region} tiene asignado en su presupuesto institucional de apertura un monto de S/ {pia_intervenciones_region}.00 \
para el financiamiento de las intervenciones y acciones pedagógicas. Asimismo, es importante mencionar que adicionalmente se \
realizará transferencias de partidas con cargo a los recursos del Ministerio de Educación a favor de los Gobierno Regionales, \
en base a los resultados de la ejecución de los recursos asignados, conforme lo dispuesto en el marco del numeral 42.1 y 42.2 del artículo 42 de la Ley de Presupuesto 2023.", style='List Bullet')
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    #interv_parrafo2.style = document.styles['Heading 5']
    #interv_parrafo2 = document.add_paragraph()

    interv_parrafo2 = document.add_paragraph("En el Anexo N° 1 puede hallarse una descripción de las intervenciones asignadas a esta región, así como su costo.", style='List Bullet')
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    #interv_parrafo2.style = document.styles['Heading 5'] 
 
    #######################################
    # Tabla 1 : Costo anual de las intervenciones y recursos disponibles en el PIA, por Unidad Ejecutora- La Libertad
    ########################################    
    int_titulo = document.add_paragraph()
    int_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    int_titulo_negrita = int_titulo.add_run(f"Tabla N° 01. Costo anual de las intervenciones y recursos disponibles en el PIA, por Unidad Ejecutora - {region}")
    #int_titulo_negrita.style = document.styles['List Bullet']
    #int_titulo_negrita.style = document.styles['List Bullet']
    int_titulo_negrita.bold = True    
    tabla1_interv = document.add_table(tabla_intervenciones_formato.shape[0]+1, tabla_intervenciones_formato.shape[1])
    tabla1_interv.autofit = False
    tabla1_interv.allow_autofit = True
    tabla1_interv.style = "tabla_minedu_1"
    #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    row = tabla1_interv.rows[0].cells
    row[0].text = "Unidad Ejecutora"
    row[1].text = "Costo anual"
    row[2].text = "PIA"
    #row[3].text = "DEV."
    #row[4].text = "% DEV"
    #row[5].text = "COSTO AL MES"
    #row[4].text = "% DEV COSTO AL MES"
    ## Contenido de la tabla
    for i in range(tabla_intervenciones_formato.shape[0]):
        for j in range(tabla_intervenciones_formato.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato.values[i,j])
    
    #Mi nuevo estilo para fuente:
    fuente_stilo = document.styles.add_style('Fuentes', WD_STYLE_TYPE.PARAGRAPH)
    fuente_stilo.font.name = 'Arial (Cuerpo)'
    fuente_stilo.font.size = docx.shared.Pt(9)
    
    #Fuente:
    interv_parrafo3 = document.add_paragraph(f"Fuente: Fuente: Base SIAF al corte de {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} del 2023 e informes UPP")
    interv_parrafo3.style = fuente_stilo
    interv_parrafo3.style.font.italic = True
    #interv_parrafo3.style.font.size = docx.shared.Pt(9)

    #############################################################################################################################################################################
    # Incluimos sección 1 de intervenciones pedagógicas: Tabla N° 02. Transferencias realizadas para el financiamiento de intervenciones #
    #############################################################################################################################################################################               
    #interv_parrafo2 = document.add_paragraph()
    
    if id_trans=="1":
        interv_parrafo2 = document.add_paragraph(f"El Gobierno Regional de {region} ha recibido un monto S/ {tr_intervenciones_region}.00 en transferencias con cargo a los recursos del Ministerio\
 de Educación a favor de los Gobierno Regionales, en base a los resultados de la ejecución de los recursos asignados, conforme lo dispuesto en el marco del numeral 42.1\
 y 42.2 del artículo 42 de la Ley de Presupuesto 2023. ", style='List Bullet')
        interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY        
        #interv_parrafo2.style = document.styles['Heading 5']
        #interv_parrafo2 = document.add_paragraph()
    
        ######################################################################################################
        # Tabla 2 : Tabla N° 02. Transferencias realizadas para el financiamiento de intervenciones
        ######################################################################################################
        int_titulo = document.add_paragraph()
        int_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        int_titulo_negrita = int_titulo.add_run(f"Tabla N° 02. Transferencias realizadas para el financiamiento de intervenciones - {region}")
        #int_titulo_negrita.style = document.styles['List Bullet']
        #int_titulo_negrita.style = document.styles['List Bullet']
        int_titulo_negrita.bold = True    
        tabla1_interv = document.add_table(tabla_intervenciones_formato_2.shape[0]+1, tabla_intervenciones_formato_2.shape[1])
        tabla1_interv.autofit = False
        tabla1_interv.allow_autofit = True
        tabla1_interv.style = "tabla_minedu_1"
        #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        row = tabla1_interv.rows[0].cells
        row[0].text = "Unidad Ejecutora"
        row[1].text = "PIM"
        row[2].text = "Primera transferencia (DS N° XX-2023-EF)"
        row[3].text = "Segunda transferencia (DS N° XX-2023-EF)"
        row[3].text = "Total transferido"
        #row[4].text = "% DEV"
        #row[5].text = "COSTO AL MES"
        #row[4].text = "% DEV COSTO AL MES"
        ## Contenido de la tabla
        for i in range(tabla_intervenciones_formato_2.shape[0]):
            for j in range(tabla_intervenciones_formato_2.shape[-1]):
                tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_2.values[i,j])
    
        #Fuente:
        interv_parrafo3 = document.add_paragraph(f"Fuente: Fuente: Base SIAF al corte de {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} del 2023 e informes UPP")
        interv_parrafo3.style = fuente_stilo
        interv_parrafo3.style.font.italic = True
        #interv_parrafo3.style.font.size = docx.shared.Pt(9)
    
    elif id_trans=="2":
        interv_parrafo2 = document.add_paragraph(f"El Gobierno Regional de {region} ha recibido un monto S/ {tr_intervenciones_region}.00 en transferencias con cargo a los recursos del Ministerio\
 de Educación a favor de los Gobierno Regionales, en base a los resultados de la ejecución de los recursos asignados, conforme lo dispuesto en el marco del numeral 42.1\
 y 42.2 del artículo 42 de la Ley de Presupuesto 2023. ", style='List Bullet')
        interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY        
        #interv_parrafo2.style = document.styles['Heading 5']
        #interv_parrafo2 = document.add_paragraph()
    
        ######################################################################################################
        # Tabla 2 : Tabla N° 02. Transferencias realizadas para el financiamiento de intervenciones
        ######################################################################################################
        int_titulo = document.add_paragraph()
        int_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        int_titulo_negrita = int_titulo.add_run(f"Tabla N° 02. Transferencias realizadas para el financiamiento de intervenciones - {region}")
        #int_titulo_negrita.style = document.styles['List Bullet']
        #int_titulo_negrita.style = document.styles['List Bullet']
        int_titulo_negrita.bold = True    
        tabla1_interv = document.add_table(tabla_intervenciones_formato_2.shape[0]+1, tabla_intervenciones_formato_2.shape[1])
        tabla1_interv.autofit = False
        tabla1_interv.allow_autofit = True
        tabla1_interv.style = "tabla_minedu_1"
        #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        row = tabla1_interv.rows[0].cells
        row[0].text = "Unidad Ejecutora"
        row[1].text = "PIM"
        row[2].text = "Primera transferencia (DS-N°031-2023-EF)"
        #row[3].text = "Segunda transferencia (DS N° XX-2023-EF)"
        row[3].text = "Total transferido"
        #row[4].text = "% DEV"
        #row[5].text = "COSTO AL MES"
        #row[4].text = "% DEV COSTO AL MES"
        ## Contenido de la tabla
        for i in range(tabla_intervenciones_formato_2.shape[0]):
            for j in range(tabla_intervenciones_formato_2.shape[-1]):
                tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_2.values[i,j])
    
        #Fuente:
        interv_parrafo3 = document.add_paragraph(f"Fuente: Fuente: Base SIAF al corte de {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} del 2023 e informes UPP")
        interv_parrafo3.style = fuente_stilo
        interv_parrafo3.style.font.italic = True
        #interv_parrafo3.style.font.size = docx.shared.Pt(9)    
    
    #############################################################################################################################################################################
    # Incluiremos informacion vinculada al grafico de ejecucion por Unidad Ejecutora                 #
    #############################################################################################################################################################################
    interv_parrafo2 = document.add_paragraph(f"Al {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} del {nyear_disponibilidad}, el Pliego Gobierno Regional de {region} ha ejecutado un monto total de S/ {ejecucion_intervenciones_region},\
 lo que representa el {porcentaje_ejecucion} del PIM.", style='List Bullet')
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY        
    #interv_parrafo2.style = document.styles['Heading 5']
    #interv_parrafo2.style = document.styles['List Bullet']
    #interv_parrafo2 = document.add_paragraph()    
    
    ##############################################################
    #Insertando el gráfico de ejecución
    #########################################################################
    int_titulo = document.add_paragraph()
    int_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    int_titulo_negrita = int_titulo.add_run("Gráfico N° 01 Avance en la ejecución presupuestal por Unidad Ejecutora (Dev/PIM%)")
    #int_titulo_negrita.style = document.styles['List Bullet']
    #int_titulo_negrita.style = document.styles['List Bullet']
    int_titulo_negrita.bold = True   
    
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    run = paragraph.add_run("")
    #run.add_picture('Gráfico_1.png', width=Inches(2))
    path_grafico_vf=path_grafico / "Gráfico_1.png"
    path_grafico_vf = path_grafico_vf.as_posix()
    run.add_picture(path_grafico_vf, width=Inches(6.0), height=Inches(3.6))
    
    #Fuente:
    interv_parrafo3 = document.add_paragraph(f"Fuente: Base SIAF al corte de {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} del 2023")
    interv_parrafo3.style = fuente_stilo
    interv_parrafo3.style.font.italic = True
    #interv_parrafo3.style.font.size = docx.shared.Pt(9)
    
    #############################################################################################################################################################################
    # Incluiremos Recursos asignados a nivel de las intervenciones y acciones pedagogicas                #
    #############################################################################################################################################################################
    interv_parrafo2 = document.add_paragraph(f"El Pliego Gobierno Regional de {region} implementa un total de {total_intervencion} de intervenciones y acciones pedagógicas para el Año 2023.\
 Las cinco intervenciones que han recibido mayores recursos en el PIM al {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} del {nyear_disponibilidad} son: {total_interven_mayor.iloc[0,0]} ({total_interven_mayor_1} millones), {total_interven_mayor.iloc[1,0]} ({total_interven_mayor_2} millones),\
 {total_interven_mayor.iloc[2,0]} ({total_interven_mayor_3} millones), {total_interven_mayor.iloc[3,0]} ({total_interven_mayor_4} millones) y {total_interven_mayor.iloc[4,0]} ({total_interven_mayor_5} millones)", style='List Bullet')
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY        
    #interv_parrafo2.style = document.styles['Heading 5']
    #interv_parrafo2.style = document.styles['List Bullet']
    #interv_parrafo2 = document.add_paragraph()    
     
    #######################################
    # Tabla 1 : Costo anual de las intervenciones y recursos disponibles en el PIA, por Unidad Ejecutora- La Libertad
    ########################################    
    int_titulo = document.add_paragraph()
    int_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    int_titulo_negrita = int_titulo.add_run(f"Tabla N° 03 Recursos asignados a nivel de las intervenciones y acciones pedagógicas")
    #int_titulo_negrita.style = document.styles['List Bullet']
    #int_titulo_negrita.style = document.styles['List Bullet']
    int_titulo_negrita.bold = True    
    tabla1_interv = document.add_table(tabla_intervenciones_formato.shape[0]+1, tabla_intervenciones_formato.shape[1])
    tabla1_interv.autofit = False
    tabla1_interv.allow_autofit = True
    tabla1_interv.style = "tabla_minedu_1"
    #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    row = tabla1_interv.rows[0].cells
    row[0].text = "Intervención pedagógica"
    row[1].text = "Costo anual"
    row[2].text = "PIA"
    #row[3].text = "DEV."
    #row[4].text = "% DEV"
    #row[5].text = "COSTO AL MES"
    #row[4].text = "% DEV COSTO AL MES"
    ## Contenido de la tabla
    for i in range(tabla_intervenciones_formato.shape[0]):
        for j in range(tabla_intervenciones_formato.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato.values[i,j])
    '''
    #Mi nuevo estilo para fuente:
    fuente_stilo = document.styles.add_style('Fuentes', WD_STYLE_TYPE.PARAGRAPH)
    fuente_stilo.font.name = 'Arial (Cuerpo)'
    fuente_stilo.font.size = docx.shared.Pt(9)
    '''
    #Fuente:
    interv_parrafo3 = document.add_paragraph(f"Fuente: Fuente: Base SIAF al corte de {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} del 2023 e informes UPP")
    interv_parrafo3.style = fuente_stilo
    interv_parrafo3.style.font.italic = True
    #interv_parrafo3.style.font.size = docx.shared.Pt(9)
    
    #############################################################################################################################################################################
    # Incluiremos Avance en Ejecución Presupuestal a nivel de intervenciones pedagogicas  tabla 04              #
    #############################################################################################################################################################################
    interv_parrafo2 = document.add_paragraph(f"La ejecución por intervención pedagógica muestra mayor avance en: {nombres_inte}. Es importante resaltar que las intervenciones pedagógicas: {nombres_inte_cero}.\
 Cuentan con recursos, pero no presentan avances.", style='List Bullet')
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY        
    #interv_parrafo2.style = document.styles['Heading 5']
    #interv_parrafo2.style = document.styles['List Bullet']
    #interv_parrafo2 = document.add_paragraph()  
    
    #######################################
    # Tabla 4 : Avance en Ejecución Presupuestal a nivel de intervenciones pedagogicas (expresados en miles de soles)
    ########################################    
    int_titulo = document.add_paragraph()
    int_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    int_titulo_negrita = int_titulo.add_run(f"Tabla N° 04 Avance en Ejecución Presupuestal a nivel de intervenciones pedagógicas (expresados en miles de soles)")
    #int_titulo_negrita.style = document.styles['List Bullet']
    #int_titulo_negrita.style = document.styles['List Bullet']
    int_titulo_negrita.bold = True    
    tabla1_interv = document.add_table(tabla_intervenciones_formato_4.shape[0]+1, tabla_intervenciones_formato_4.shape[1])
    tabla1_interv.autofit = False
    tabla1_interv.allow_autofit = True
    tabla1_interv.style = "tabla_minedu_1"
    #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    row = tabla1_interv.rows[0].cells
    row[0].text = "Intervención pedagógica"
    row[1].text = "PIM"
    row[2].text = "Devengado"
    row[3].text = "% Ejec"
    #row[4].text = "% DEV"
    #row[5].text = "COSTO AL MES"
    #row[4].text = "% DEV COSTO AL MES"
    ## Contenido de la tabla
    for i in range(tabla_intervenciones_formato_4.shape[0]):
        for j in range(tabla_intervenciones_formato_4.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_4.values[i,j])  
    
    #Fuente:
    interv_parrafo3 = document.add_paragraph(f"Fuente: Base SIAF al corte de {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} del 2023 e informes UPP")
    interv_parrafo3.style = fuente_stilo
    interv_parrafo3.style.font.italic = True
    #interv_parrafo3.style.font.size = docx.shared.Pt(9)
    
    #######################################
    # Tabla 5 : PEAS asignadas y costo anual de contratación- La Libertad
    ########################################  
    # Texto
    document.add_page_break()
    document.add_heading("2. Contratación de Personal CAS para IAP", level=1) # 1) Intervenciones pedagógicas  
    #document.add_heading(f"Corte: {fecha_corte_disponibilidad_format}", level=3)
    
    interv_parrafo2 = document.add_paragraph(f"El Pliego Gobierno Regional de {region} tiene un total de {total_peas_region} PEAS Programadas,\
 para las {cant_ue_region} Unidades Ejecutoras que la conforman. En total, el costo estimado anual de las contrataciones es de S/ {costo_anual_region} millones de soles.", style='List Bullet')
    interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY        
    #interv_parrafo2.style = document.styles['Heading 5']
    #interv_parrafo2.style = document.styles['List Bullet']
    #interv_parrafo2 = document.add_paragraph()  
    
    # Tabla
    int_titulo = document.add_paragraph()
    int_titulo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    int_titulo_negrita = int_titulo.add_run(f"Tabla N° 05 PEAS asignadas y costo anual de contratación- {region}")
    #int_titulo_negrita.style = document.styles['List Bullet']
    #int_titulo_negrita.style = document.styles['List Bullet']
    int_titulo_negrita.bold = True    
    tabla1_interv = document.add_table(tabla_intervenciones_formato_6.shape[0]+1, tabla_intervenciones_formato_6.shape[1])
    tabla1_interv.autofit = False
    tabla1_interv.allow_autofit = True
    tabla1_interv.style = "tabla_minedu_1"
    #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    row = tabla1_interv.rows[0].cells
    row[0].text = "Unidad Ejecutora"
    row[1].text = "PEAS programadas"
    row[2].text = "Costo anual estimado"
    #row[3].text = "% Ejec"
    #row[4].text = "% DEV"
    #row[5].text = "COSTO AL MES"
    #row[4].text = "% DEV COSTO AL MES"
    ## Contenido de la tabla
    for i in range(tabla_intervenciones_formato_6.shape[0]):
        for j in range(tabla_intervenciones_formato_6.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_6.values[i,j])  
    
    #Fuente:
    interv_parrafo3 = document.add_paragraph(f"Fuente: Base NEXUS al corte de {ndia_nexus} de {nmes_nex[nmes_nex_entero-1]} del 2023 e informes UPP")
    interv_parrafo3.style = fuente_stilo
    interv_parrafo3.style.font.italic = True
    #interv_parrafo3.style.font.size = docx.shared.Pt(9)
    
    #######################
    # Guardamos documento #
    #######################
    document.save(nueva_carpeta / f'AM_{region}_{fecha_actual}.docx')
    
    #Inseraremos inforamcion de algunas variables que tenemos en word
    #--------------------------------------------------------------------
    #El nombre de region en la caratural, los resumenes de variables y el gráfico de mapa
    doc = DocxTemplate(nueva_carpeta / f'AM_{region}_{fecha_actual}.docx')
    
    path_mapa=path_mapas / f"{region}.png"
    path_mapa = path_mapa.as_posix()

    mi_region=region
    mi_ejecutora=cant_ue_region
    cost_interven=pia_intervenciones_region
    fecha_siaf=fecha_corte_disponibilidad_format
    fecha_nexus=fecha_corte_nexus_format
    transf_materia=transf_materiales
    
    context = {'mi_region': mi_region, 'mi_ejecutora': mi_ejecutora, 'cost_interven':cost_interven,
               'fecha_siaf':fecha_siaf, 'fecha_nexus':fecha_nexus, 'mi_imagen':InlineImage(doc, path_mapa, width=Cm(7.38), height=Cm(9.82)),
               'transf_materia': transf_materia}
    
    doc.render(context)
    doc.save(nueva_carpeta / f'AM_{region}_{fecha_actual}.docx')
    
    #Finalmente cargaremos el word para actualizar el indice 
    #--------------------------------------------------------------------
    
    # Crear una instancia de la aplicación Word
    word = win32.Dispatch('Word.Application')

    # Abrir el archivo de Word y no activamos su ventana
    file_path = nueva_carpeta / f'AM_{region}_{fecha_actual}.docx'
    file_path_2 = file_path.as_posix()
    file_path_3 = file_path_2.replace('/', '\\')
    
    doc = word.Documents.Open(file_path_3)
    word.Visible = False

    # Actualizar la tabla de contenido (Como es la primera tabla de contenidos estou utilizando "1")
    doc.TablesOfContents(1).Update()

    # La seccion que en blanco es la seccion 4
    #section_range = doc.Range(doc.Sections(4).Range.Start, doc.Sections(4).Range.End)
    #section_range.Delete() # eliminar la sección del documento

    doc.Close(SaveChanges=True) # guardar y cerrar el documento
    #word.Quit() # cerrar la aplicación de Word

    # Guardar y cerrar el archivo
    #file_path=r'C:\Users\ANALISTAUP29\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada\Am prueba v4.docx'
    #doc.Save(file_path)
    #doc.Close()

    # Cerrar la aplicación Word
    word.Quit()

###########################################################
# Creamos tabla con lista de files para enviar por correo #
###########################################################
    
# Generamos lista de AM.
#lista_AM = glob.glob(os.path.join(proyecto, f"output/AM_corta_region/AM_{fecha_actual}/*"))

#lista_regiones = pd.DataFrame (lista_AM)
#lista_regiones.rename( columns={0:'path'}, inplace=True )
#lista_regiones[['a', 'b', 'c']] = lista_regiones["path"].str.split("AM_", expand = True)
#lista_regiones[['date', 'e']] = lista_regiones["b"].str.split("/", expand = True)
#lista_regiones[['region', 'g']] = lista_regiones["c"].str.split("_", expand = True)
#lista_regiones = lista_regiones[["path", "date","region"]]
#lista_regiones.to_excel(Path(proyecto, "documentacion", "lista_regiones.xlsx"), index = False)
